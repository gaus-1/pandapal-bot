from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading_center(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def main():
    target_docx = (
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal).docx"
    )

    print("Opening DOCX for diagram injection...")
    doc = Document(target_docx)

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ А")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Схема базы данных (ERD-диаграмма)")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\erd.png", width=Inches(6.0)
    )

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ Б")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Архитектура программного комплекса")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\arch.png", width=Inches(6.0)
    )

    doc.save(target_docx)
    print("Successfully added ERD and Arch diagrams to DOCX!")


if __name__ == "__main__":
    main()
