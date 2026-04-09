import docx

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(doc_path)

with open(r"c:\Users\Vyacheslav\PandaPal\tmp_verify.txt", "w", encoding="utf-8") as f:
    f.write("--- FIRST 15 PARAGRAPHS ---\n")
    for i in range(min(15, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        f.write(
            f"[{i}] len={len(p.text)} align={p.alignment} indent={p.paragraph_format.first_line_indent} TEXT: '{p.text}'\n"
        )

    f.write("\n--- TABLE OF CONTENTS (TOC) Sample ---\n")
    # find lines with a tab stop
    toc_lines = 0
    for p in doc.paragraphs:
        if "\t" in p.text and (
            p.text.endswith("5")
            or p.text.endswith("6")
            or p.text.strip().rsplit("\t", 1)[-1].isdigit()
        ):
            f.write(f"TOC line: align={p.alignment} TEXT: '{p.text}'\n")
            toc_lines += 1
            if toc_lines > 5:
                break

    f.write("\n--- SCHEMA CAPTIONS ---\n")
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Рисунок"):
            f.write(
                f"[{i}] align={p.alignment} indent={p.paragraph_format.first_line_indent} TEXT: '{p.text}'\n"
            )

    f.write("\n--- DIAGRAM VALIDATION ---\n")
    # Finding ER diagram
    er_found = False
    for i, p in enumerate(doc.paragraphs):
        if "telegram_id (PK)" in p.text:
            er_found = True
            font_name = "None"
            if p.runs:
                font_name = p.runs[0].font.name
            f.write(f"ER diagram box line [{i}]: font={font_name} text='{p.text}'\n")

    if not er_found:
        f.write("ERROR: ER DIAGRAM NOT FOUND!\n")

    f.write(f"\nTotal paragraphs: {len(doc.paragraphs)}\n")
