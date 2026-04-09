import docx

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(doc_path)
with open(r"c:\Users\Vyacheslav\PandaPal\tmp_verify2.txt", "w", encoding="utf-8") as f:
    for i in range(len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text.startswith("Рисунок"):
            # print surrounding 2 paragraphs
            for offset in [-2, -1, 0, 1]:
                if 0 <= i + offset < len(doc.paragraphs):
                    f.write(f"[{i+offset}] {doc.paragraphs[i+offset].text}\n")
            f.write("-" * 20 + "\n")
