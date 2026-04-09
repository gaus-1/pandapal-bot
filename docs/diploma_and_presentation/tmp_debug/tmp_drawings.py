import docx

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
doc = docx.Document(doc_path)

drawings = []
for i, p in enumerate(doc.paragraphs):
    if "w:drawing" in p._p.xml:
        drawings.append(i)

with open(r"c:\Users\Vyacheslav\PandaPal\tmp_drawings.txt", "w", encoding="utf-8") as f:
    f.write(f"Paragraphs with drawings: {drawings}\n")
    for d in drawings:
        p_text = doc.paragraphs[d].text
        next_text = doc.paragraphs[d + 1].text if d + 1 < len(doc.paragraphs) else ""
        f.write(f"P[{d}] text: {p_text} | Next[{d+1}] text: {next_text}\n")
