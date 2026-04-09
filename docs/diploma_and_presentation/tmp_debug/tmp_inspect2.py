import docx

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
try:
    doc = docx.Document(doc_path)

    with open(r"c:\Users\Vyacheslav\PandaPal\tmp_inspect.txt", "w", encoding="utf-8") as f:
        f.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Total tables: {len(doc.tables)}\n")
        f.write(f"Total inline shapes: {len(doc.inline_shapes)}\n")

        f.write("\n--- First 50 Paragraphs ---\n")
        for i, p in enumerate(doc.paragraphs[:50]):
            f.write(f"[{i}] {p.text}\n")

        f.write("\n--- Table Headers ---\n")
        for i, t in enumerate(doc.tables):
            cells = t.rows[0].cells
            header_text = " | ".join([c.text.replace("\n", " ").strip() for c in cells])
            f.write(f"Table {i+1}: {header_text}\n")

except Exception as e:
    import traceback

    with open(r"c:\Users\Vyacheslav\PandaPal\tmp_inspect.txt", "w", encoding="utf-8") as f:
        f.write(f"Error: {e}\n{traceback.format_exc()}")
