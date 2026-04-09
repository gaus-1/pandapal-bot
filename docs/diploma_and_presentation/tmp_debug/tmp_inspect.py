import docx

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
try:
    doc = docx.Document(doc_path)

    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total inline shapes: {len(doc.inline_shapes)}")

    print("\n--- First 20 Paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"[{i}] {p.text}")

except Exception as e:
    print(f"Error: {e}")
