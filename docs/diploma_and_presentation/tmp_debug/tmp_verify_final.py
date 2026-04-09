"""Quick verification of formatting results."""

import logging

import docx
from docx.oxml.ns import qn

logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(doc_path)

out = r"c:\Users\Vyacheslav\PandaPal\tmp_verify_final.txt"

with open(out, "w", encoding="utf-8") as f:
    f.write("=== VERIFICATION REPORT ===\n\n")

    # 1. Check centered items
    centered_count = 0
    not_centered_diagrams = []
    not_centered_captions = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        align = str(p.alignment)

        # Diagrams (box-drawing)
        has_box = any(c in text for c in "┌│└├─┐┘┤┬┴┼▼▲╔╗╚╝╠╣╦╩╬═║")
        if has_box:
            if "CENTER" in align:
                centered_count += 1
            else:
                not_centered_diagrams.append(f"  [{i}] align={align} | {text[:60]}")

        # Figure captions
        if text.startswith("Рисунок") and "—" in text:
            if "CENTER" not in align:
                not_centered_captions.append(f"  [{i}] align={align} | {text}")

        # Table captions
        if text.startswith("Таблица") and "—" in text:
            if "CENTER" not in align:
                not_centered_captions.append(f"  [{i}] align={align} | {text}")

    f.write(f"Centered diagram paragraphs: {centered_count}\n")
    if not_centered_diagrams:
        f.write(f"NOT centered diagrams ({len(not_centered_diagrams)}):\n")
        for d in not_centered_diagrams[:10]:
            f.write(d + "\n")
    else:
        f.write("✓ All diagram paragraphs are centered!\n")

    if not_centered_captions:
        f.write(f"\nNOT centered captions ({len(not_centered_captions)}):\n")
        for c in not_centered_captions:
            f.write(c + "\n")
    else:
        f.write("✓ All captions (Рисунок/Таблица) are centered!\n")

    # 2. Tables alignment
    f.write("\n--- TABLE ALIGNMENT ---\n")
    for i, table in enumerate(doc.tables):
        tblPr = table._tbl.find(qn("w:tblPr"))
        jc_val = "unknown"
        if tblPr is not None:
            jc = tblPr.find(qn("w:jc"))
            if jc is not None:
                jc_val = jc.get(qn("w:val"), "unknown")
        f.write(f"Table {i}: alignment={jc_val}\n")

    # 3. ER Diagram check
    f.write("\n--- ER DIAGRAM CHECK ---\n")
    for i, p in enumerate(doc.paragraphs):
        if "СХЕМА БАЗЫ ДАННЫХ (ERD)" in p.text:
            f.write(f"Found ER header at [{i}]: {p.text[:60]}\n")
            f.write(f"  alignment: {p.alignment}\n")
            # Print next 10 lines
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                f.write(f"  [{j}] {doc.paragraphs[j].text[:80]}\n")
            break

    f.write("\n✅ Verification complete\n")

logger.info(f"Written to {out}")
