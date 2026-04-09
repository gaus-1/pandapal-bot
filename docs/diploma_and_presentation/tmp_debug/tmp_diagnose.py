"""Diagnostic script: dump full structure of the final docx."""

import logging

import docx
from docx.oxml.ns import qn

logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(doc_path)

out = r"c:\Users\Vyacheslav\PandaPal\tmp_diagnose.txt"

with open(out, "w", encoding="utf-8") as f:
    f.write("=" * 80 + "\n")
    f.write("DOCUMENT STRUCTURE ANALYSIS\n")
    f.write("=" * 80 + "\n\n")

    # 1. Sections info
    f.write("--- SECTIONS ---\n")
    for i, section in enumerate(doc.sections):
        f.write(
            f"Section {i}: page_width={section.page_width}, page_height={section.page_height}\n"
        )
        f.write(
            f"  margins: left={section.left_margin}, right={section.right_margin}, top={section.top_margin}, bottom={section.bottom_margin}\n"
        )
    f.write("\n")

    # 2. Tables info
    f.write("--- TABLES ---\n")
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0
        # Get alignment
        tbl_xml = table._tbl
        jc = tbl_xml.find(qn("w:tblPr"))
        alignment = "unknown"
        if jc is not None:
            jc_elem = jc.find(qn("w:jc"))
            if jc_elem is not None:
                alignment = jc_elem.get(qn("w:val"), "unknown")

        # First cell text for identification
        first_cell = ""
        if rows > 0 and cols > 0:
            first_cell = table.cell(0, 0).text[:50].replace("\n", " ")

        f.write(f"Table {i}: {rows}x{cols}, alignment={alignment}, first_cell='{first_cell}'\n")
    f.write("\n")

    # 3. All paragraphs with key info
    f.write("--- PARAGRAPHS ---\n")
    for i, p in enumerate(doc.paragraphs):
        text = p.text[:120].replace("\n", "\\n")
        alignment = str(p.alignment) if p.alignment else "None"
        style = p.style.name if p.style else "None"

        # Check if paragraph has images
        has_image = False
        for run in p.runs:
            if run._r.findall(qn("w:drawing")) or run._r.findall(qn("w:pict")):
                has_image = True
                break

        # Check for box-drawing chars
        has_box = any(c in p.text for c in "┌│└├─┐┘┤┬┴┼▼▲►◄═║╔╗╚╝╠╣╦╩╬")

        indent = p.paragraph_format.first_line_indent
        left_indent = p.paragraph_format.left_indent

        markers = []
        if has_image:
            markers.append("IMG")
        if has_box:
            markers.append("BOX")
        if text.strip().startswith("Рисунок"):
            markers.append("FIG_CAPTION")
        if text.strip().startswith("Таблица"):
            markers.append("TBL_CAPTION")
        if text.strip().startswith("Приложение"):
            markers.append("APPENDIX")
        if "Схема" in text or "схема" in text:
            markers.append("SCHEMA")

        marker_str = f" [{', '.join(markers)}]" if markers else ""
        f.write(f"[{i:4d}] align={alignment:30s} style={style:20s}{marker_str} | {text}\n")

    # 4. Check body elements for drawings/images outside paragraphs
    f.write("\n--- BODY ELEMENTS WITH IMAGES ---\n")
    body = doc.element.body
    for i, elem in enumerate(body):
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        drawings = elem.findall(".//" + qn("w:drawing"))
        picts = elem.findall(".//" + qn("w:pict"))
        if drawings or picts:
            text_content = elem.text or ""
            # Try to get text from w:t elements
            t_elems = elem.findall(".//" + qn("w:t"))
            text_from_t = " ".join([t.text or "" for t in t_elems])[:80]
            f.write(
                f"Body[{i}] tag={tag}, drawings={len(drawings)}, picts={len(picts)}, text='{text_from_t}'\n"
            )

    f.write("\n--- INLINE SHAPES ---\n")
    for i, shape in enumerate(doc.inline_shapes):
        f.write(f"InlineShape {i}: type={shape.type}, width={shape.width}, height={shape.height}\n")
        # Find parent paragraph
        parent = shape._inline.getparent().getparent()
        # Find paragraph index
        for j, p in enumerate(doc.paragraphs):
            if p._p is parent:
                f.write(f"  -> in paragraph [{j}]: '{p.text[:60]}'\n")
                f.write(f"  -> paragraph alignment: {p.alignment}\n")
                break

logger.info(f"Diagnostic output written to {out}")
