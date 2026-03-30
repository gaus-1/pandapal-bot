import base64

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# 1. ERD
erd_mermaid = """
erDiagram
    USERS ||--o{ PANDA_PET : "has"
    USERS ||--o{ PAYMENTS : "makes"
    USERS ||--o{ CHAT_MESSAGES : "sends"
    USERS {
        int id PK
        bigint telegram_id
        string current_state
        boolean is_premium
        int total_coins
    }
    PANDA_PET {
        int id PK
        int user_id FK
        int health
        int energy
        int satiety
        int mood
    }
    CHAT_MESSAGES {
        int id PK
        int user_id FK
        text role
        text content
        timestamp created_at
    }
    PAYMENTS {
        int id PK
        int user_id FK
        string yookassa_id
        string status
        decimal amount
    }
"""

# 2. Architecture
arch_mermaid = """
graph TB
    subgraph Frontend [Клиентская часть]
        TG[Telegram Client]
        MAPP[Mini App React 19]
        WEB[Web Landing]
    end

    subgraph Backend [Backend aiohttp + aiogram]
        BOT[Tg Bot Webhook]
        API[Fast API / aiohttp]
        MOD[4-Layer AI Moderation]
        RAG[RAG Engine]
        GAME[Game & Pet Engine]
    end

    subgraph Data [Базы данных]
        PG[(PostgreSQL + pgvector)]
        REDIS[(Redis In-Memory)]
    end

    subgraph External [Внешние API]
        YGPT[YandexGPT Pro]
        YVIS[Yandex Vision OCR]
        YOO[YooKassa Webhooks]
    end

    TG --> BOT
    MAPP --> API
    WEB --> API
    BOT --> MOD
    API --> MOD
    MOD --> RAG
    MOD --> GAME
    RAG --> YGPT
    RAG --> PG
    GAME --> REDIS
    API --> YVIS
    API -.-> YOO
"""

# 3. Moderation Flow
mod_mermaid = """
sequenceDiagram
    participant U as User (Школьник)
    participant B as Backend
    participant M as AI Moderation
    participant DB as Vector DB
    participant LLM as YandexGPT

    U->>B: Задает вопрос (с тетрадью)
    B->>M: RegExp + Levenshtein (Оскорбления)
    alt Мат найден
        M-->>U: Блокировка!
    else Чисто
        M->>M: ML Классификация намерений
        alt Вредное намерение
            M-->>U: Блокировка темы
        else Образовательный контекст
            M->>DB: Векторизация и поиск фактов
            DB-->>M: Найденный параграф
            M->>LLM: Контекст + Запрос
            LLM-->>U: SSE Поток с ответом
        end
    end
"""


def get_mermaid_url(code):
    encoded = base64.urlsafe_b64encode(code.encode("utf8")).decode("utf8")
    return f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"


def download_image(code, filename):
    print(f"Downloading {filename}...")
    url = get_mermaid_url(code)
    r = requests.get(url)
    with open(filename, "wb") as f:
        f.write(r.content)


def add_heading_center(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def main():
    target_docx = (
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal).docx"
    )

    download_image(erd_mermaid, "erd.png")
    download_image(arch_mermaid, "arch.png")
    download_image(mod_mermaid, "mod.png")

    doc = Document(target_docx)

    # Add page break
    doc.add_page_break()

    # Appendices
    add_heading_center(doc, "ПРИЛОЖЕНИЕ А")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Схема базы данных (ERD-диаграмма)")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("erd.png", width=Inches(5))

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ Б")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Архитектура программного комплекса")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("arch.png", width=Inches(6))

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ В")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Sequence-диаграмма работы 4-уровневого ИИ-модератора")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("mod.png", width=Inches(5.0))

    # Save overwriting
    doc.save(target_docx)
    print("Successfully added 3 diagrams to DOCX Appendices!")


if __name__ == "__main__":
    main()
