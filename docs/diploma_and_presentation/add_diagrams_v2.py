import base64
import sys

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# 1. ERD
erd_mermaid = """
erDiagram
    USERS ||--o{ PANDA_PET : "has"
    USERS ||--o{ PAYMENTS : "makes"
    USERS ||--o{ CHAT_MESSAGES : "sends"
    USERS {
        int id PK
        bigint telegram_id
        boolean premium
        int coins
    }
    PANDA_PET {
        int id PK
        int user_id FK
        int health
        int mood
    }
    CHAT_MESSAGES {
        int id PK
        int user_id FK
        text role
        text content
    }
    PAYMENTS {
        int id PK
        int user_id FK
        string status
        decimal amount
    }
"""


def get_mermaid_url(code):
    encoded = base64.b64encode(code.encode("utf8")).decode("utf8")
    return f"https://mermaid.ink/img/{encoded}?type=png"


def download_image(code, filename):
    print(f"Downloading {filename}...")
    url = get_mermaid_url(code)
    try:
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        with open(filename, "wb") as f:
            f.write(r.content)
        print(f"Success: {filename} saved.")
    except Exception as e:
        print(f"Failed to download {filename}: {e}")
        sys.exit(1)


def add_heading_center(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def main():
    target_docx = (
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal).docx"
    )

    download_image(erd_mermaid, "erd.png")

    print("Opening DOCX...")
    doc = Document(target_docx)

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ А")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Схема базы данных (ERD-диаграмма)")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture("erd.png", width=Inches(6.0))

    doc.save(target_docx)
    print("Successfully added ERD diagram to DOCX Appendices!")


if __name__ == "__main__":
    main()
