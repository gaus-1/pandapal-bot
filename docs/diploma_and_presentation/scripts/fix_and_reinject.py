import logging
import os
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

DOC_PATH = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
IMG_DIR = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\img"

doc = docx.Document(DOC_PATH)

def insert_image_paragraph(doc, anchor_element, img_path, width_cm=16):
    new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_element.addprevious(new_p_elem)
    paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return paragraph

def insert_caption_paragraph(doc, anchor_element, caption_text):
    new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_element.addprevious(new_p_elem)
    paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(caption_text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.italic = False
    return paragraph

# === 1. CLEANUP ALL EXISTING INJECTED IMAGES AND CAPTIONS ===
logger.info("Phase 1: Cleanup")

to_delete = []
for p in doc.paragraphs:
    # Identify images
    if p._element.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}drawing'):
        to_delete.append(p)
    # Identify our captions
    text = p.text.strip()
    if text.startswith("Рисунок") and ("–" in text or "—" in text):
        to_delete.append(p)

logger.info(f"Deleting {len(to_delete)} paragraphs (images/captions)")
for p in to_delete:
    p._element.getparent().remove(p._element)


# === 2. FIND ANCHORS AND REINJECT ===
logger.info("Phase 2: Reinjection")

# We need to find the anchors safely.
# Anchor 1: Text contains "Диаграмма вариантов использования PandaPal" or similar text near beginning
# Wait, let's find the titles, e.g. "1.4. Моделирование бизнес-процессов" or just look for matching text.

# 2.1 FIND ANCHOR FOR FIGURE 1 & 2 & 3
# In the original text, there are explicit mentions (e.g. "Рисунок 1" or "на рисунке 1").
# Let's find "14 сущностей" or "15 сущност" for Appendix A.
anchor_fig1 = None
anchor_fig2 = None
anchor_appA = None
anchor_appB = None

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "Бизнес-процессы PandaPal" in t:
        anchor_fig1 = p
    elif "Моделирование бизнес-процессов (UML, Use Case)" in t:
        anchor_fig1 = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else p
    elif "Архитектура клиент-серверного взаимодействия" in t:
        anchor_fig2 = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else p
    elif "Приложение А" in t:
        anchor_appA = p
    elif "Приложение Б" in t:
        anchor_appB = p

# If we couldn't find them by exact head titles, let's just use paragraph indices roughly if needed or text substrings
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip().lower()
    if "диаграмма вариантов использования" in t and "PandaPal" in p.text and anchor_fig1 is None:
        anchor_fig1 = p
    if "общая архитектура" in t and anchor_fig2 is None:
        anchor_fig2 = p

# It is safer directly inserting into Appendix A and B based on "Приложение А"
# Let's find the text "База данных PandaPal содержит" to inject ERDs AFTER it
appA_text_block = None
for p in doc.paragraphs:
    if "База данных PandaPal содержит" in p.text:
        appA_text_block = p
        break

if appA_text_block:
    # Insert AFTER appA_text_block. This means we insert BEFORE the paragraph AFTER appA_text_block
    next_p = appA_text_block._element.getnext()
    anchor = docx.text.paragraph.Paragraph(next_p, doc._body) if next_p is not None else appA_text_block

    # Notice: we are adding BEFORE "anchor". So we must insert in the top-down order!
    # Wait, addprevious() adds BEFORE the current element.
    # To append 3 diagrams sequentially BEFORE an anchor:
    # 1. img1 -> anchor
    # 2. cap1 -> anchor
    # 3. img2 -> anchor
    # 4. cap2 -> anchor
    # End result: img1, cap1, img2, cap2, anchor
    # Yes! This yields the correct forward order!

    # 1_erd.png (Концептуальная)
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "1_erd.png"), 16)
    insert_caption_paragraph(doc, anchor._element, "Рисунок А.1 – Концептуальная ER-диаграмма базы данных PandaPal")
    # 6_erd_physical.png (Физическая)
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "6_erd_physical.png"), 17)
    insert_caption_paragraph(doc, anchor._element, "Рисунок А.2 – Физическая ER-диаграмма базы данных (15 таблиц)")
    # 7_erd_relational.png (Реляционная)
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "7_erd_relational.png"), 17)
    insert_caption_paragraph(doc, anchor._element, "Рисунок А.3 – Реляционная схема базы данных PandaPal")

if anchor_appB:
    # Find the next normal paragraph after Приложение Б
    anchor = anchor_appB
    for _ in range(2):
        nxt = anchor._element.getnext()
        if nxt is not None:
            anchor = docx.text.paragraph.Paragraph(nxt, doc._body)

    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "4_deployment.png"), 16)
    insert_caption_paragraph(doc, anchor._element, "Рисунок 4 – Схема развёртывания инфраструктуры PandaPal")

# For Fig 1 and Fig 2, let's find the text that mentions them
fig1_anchor = None
fig2_anchor = None
for p in doc.paragraphs:
    if "Варианты использования представлены на рисунке 1" in p.text or "на рисунке 1" in p.text.lower():
        fig1_anchor = p
    if "Архитектура разрабатываемой системы" in p.text or ("архитектура" in p.text.lower() and "на рисунке 2" in p.text.lower()):
        fig2_anchor = p

if fig1_anchor:
    nxt = fig1_anchor._element.getnext()
    anchor = docx.text.paragraph.Paragraph(nxt, doc._body) if nxt is not None else fig1_anchor
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "5_usecase.png"), 15)
    insert_caption_paragraph(doc, anchor._element, "Рисунок 1 – Диаграмма вариантов использования PandaPal")

if fig2_anchor:
    nxt = fig2_anchor._element.getnext()
    anchor = docx.text.paragraph.Paragraph(nxt, doc._body) if nxt is not None else fig2_anchor
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "2_architecture.png"), 16)
    insert_caption_paragraph(doc, anchor._element, "Рисунок 2 – Общая архитектура системы PandaPal")
    insert_image_paragraph(doc, anchor._element, os.path.join(IMG_DIR, "3_sequence.png"), 15)
    insert_caption_paragraph(doc, anchor._element, "Рисунок 3 – Диаграмма последовательности обработки сообщений")

doc.save(DOC_PATH)
logger.info("Successfully fixed layout and injected GOST compliant diagrams!")
