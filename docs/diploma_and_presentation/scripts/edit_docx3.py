import docx
import re
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(doc_path)

for p in doc.paragraphs:
    # Fix TOC
    if "\t" in p.text and p.text.strip().split("\t")[-1].strip().isdigit():
        parts = p.text.split("\t")
        if len(parts) >= 2:
            title = "".join(parts[:-1]).strip()
            page = parts[-1].strip()

            # Clean trailing periods and ellipsis from title
            title = re.sub(r'[\.\…\s]+$', '', title)

            p.text = f"{title}\t{page}"

            # Tab stops
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.paragraph_format.first_line_indent = 0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Remove duplicate Schema captions if accidentally added
    # For instance if we have "Рисунок 1 - ..." followed directly by "Рисунок 2 - ..."
    # Actually, we can check if it's TOC by looking for another match.
    # What about those lines containing "………….	4"? We fixed them above.

    # What if a TOC line didn't get a tab?
    # Some lines might have space number like "ВВЕДЕНИЕ 5"
    else:
        text_val = p.text.strip()
        # Look for "Title..... 5" or "Title… 5"
        if re.search(r'[\.\…]{2,}\s*\d+$', text_val):
            # Extract number
            match = re.search(r'^(.*?)([\.\…\s]+)(\d+)$', text_val)
            if match:
                title = match.group(1).rstrip('.\… ')
                page = match.group(3)
                p.text = f"{title}\t{page}"
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.clear_all()
                tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                p.paragraph_format.first_line_indent = 0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# I will also remove duplicate or empty paragraphs from top header if there are too many
# Paragraphs 0 is the full title. Paragraphs 1,2,3 are empty?
# Wait, let's keep one empty paragraph and delete the rest from 1,2,3 if they are empty
empty_header_count = 0
paragraphs_to_remove = []
for i in range(1, min(6, len(doc.paragraphs))):
    if not doc.paragraphs[i].text.strip():
        empty_header_count += 1
        if empty_header_count > 1:
            paragraphs_to_remove.append(doc.paragraphs[i])

for p in paragraphs_to_remove:
    p._element.getparent().remove(p._element)

doc.save(doc_path)
