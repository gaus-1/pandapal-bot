import docx
import re
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
output_path = doc_path  # We'll overwrite it inline.

doc = docx.Document(doc_path)

new_er_diagram = """
                              ┌──────────────────┐
                              │      users       │
                              │──────────────────│
                              │ telegram_id (PK) │
                              └────────┬─────────┘
      ┌─────────────────┬──────────────┼──────────────┬─────────────────┐
      │                 │              │              │                 │
      ▼                 ▼              ▼              ▼                 ▼
┌───────────┐     ┌───────────┐  ┌───────────┐  ┌───────────┐     ┌───────────┐
│panda_pets │     │ referrals │  │chat_histor│  │ learning  │     │   games   │
│───────────│     │───────────│  │───────────│  │───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│  │user_id(FK)│  │user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘  └───────────┘  └───────────┘     └───────────┘
      │                 │              │              │                 │
      ▼                 ▼              ▼              ▼                 ▼
┌───────────┐     ┌───────────┐  ┌───────────┐  ┌───────────┐     ┌───────────┐
│ payments  │     │ payouts   │  │   daily   │  │ homework  │     │game_stats │
│───────────│     │───────────│  │───────────│  │───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│  │user_id(FK)│  │user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘  └───────────┘  └───────────┘     └───────────┘
      │                 │
      ▼                 ▼
┌───────────┐     ┌───────────┐
│ subs      │     │ analytics │
│───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘
"""

# -----------------
# 1. Update text "10 ORM-моделей" to "14 ORM-моделей"
# -----------------
for p in doc.paragraphs:
    if "10 ORM-моделей" in p.text:
        p.text = p.text.replace("10 ORM-моделей", "14 ORM-моделей")

# -----------------
# 2. Table of Contents alignment (TOC)
# -----------------
toc_pattern = re.compile(r'^(.*?)\s*\.{3,}\s*(\d+)$')
for p in doc.paragraphs:
    match = toc_pattern.match(p.text.strip())
    # Or maybe it has tabs/spaces and dots intermittently
    text_val = p.text
    if len(re.findall(r'\.{3,}', text_val)) > 0 and bool(re.search(r'\d+$', text_val.strip())):
        # It's a TOC line.
        # Extract title and page number
        title_match = re.match(r'^(.*?)(?:\s*\.{2,}\s*|\s+)(\d+)$', text_val.strip())
        if title_match:
            title = title_match.group(1).rstrip('.')
            page = title_match.group(2)

            p.text = f"{title}\t{page}"

            # Add Right Tab Stop with Dot Leader
            # 16.5 cm is typical for A4 with standard margins. Let's trace it.
            tab_stops = p.paragraph_format.tab_stops
            # Clear existing tabs
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.paragraph_format.first_line_indent = 0 # No first line indent in TOC
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# -----------------
# 3. Replace ER Diagram
# -----------------
# We must find where the ER diagram starts and ends.
# It starts around "┌──────────────┐" and ends at "└──────────────┘" (in Приложение А).
er_start_idx = -1
er_end_idx = -1
in_er_diagram = False

for i, p in enumerate(doc.paragraphs):
    if "Приложение А. Схема базы данных" in p.text or "База данных PandaPal содержит" in p.text:
        # The ER diagram is shortly after this
        continue

    if getattr(p, '_er_checked', False):
        continue

    if "┌──────────────┐" in p.text and "chat_messages" in p.text:
        # But maybe we already modified it? In previous step we didn't touch it.
        pass

    # Let's use a simpler heuristic: find the block of lines that contain box drawing chars mapped to the ER diagram
    # Specifically, look for lines with ┌, │, └ right after "Приложение А"

# Actually, since I know the text, I'll just find the first paragraph after "База данных PandaPal содержит 14 ORM-моделей..."
# and if it has ascii box chars, it's the start
found_appendix_a = False
for i, p in enumerate(doc.paragraphs):
    if "Приложение А." in p.text:
        found_appendix_a = True

    if found_appendix_a and any(c in p.text for c in "┌│└├─"):
        # This is the ER diagram start
        er_start_idx = i
        # Find where it ends
        for j in range(i, len(doc.paragraphs)):
            if not any(c in doc.paragraphs[j].text for c in "┌│└├─") and len(doc.paragraphs[j].text.strip()) > 0:
                er_end_idx = j - 1
                break
        break

if er_start_idx != -1 and er_end_idx != -1:
    # Remove old ER diagram paragraphs
    for j in range(er_start_idx, er_end_idx + 1):
        p = doc.paragraphs[er_start_idx]
        p._element.getparent().remove(p._element)

    # Insert new ER diagram
    # We'll split by newline and insert paragraphs where the old one was removed
    # Wait, the element we removed is gone. We need an anchor to insert BEFORE.
    # The anchor is the paragraph at er_start_idx (which is now the paragraph that originally was er_end_idx + 1)
    if er_start_idx < len(doc.paragraphs):
        anchor_p = doc.paragraphs[er_start_idx]

        for line in new_er_diagram.strip('\n').split('\n'):
            new_p = docx.oxml.OxmlElement('w:p')
            anchor_p._element.addprevious(new_p)
            para = docx.text.paragraph.Paragraph(new_p, doc._body)
            para.text = line
            # Keep monospace formatting by using a specific font if possible,
            # and left align with zero indent
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = 0
            # Set font to Courier New
            for run in para.runs:
                run.font.name = 'Courier New'

# -----------------
# 4. Center schemas (pictures) and remove empty paragraphs
# -----------------
empty_count = 0
paragraphs_to_remove = []

for p in doc.paragraphs:
    text = p.text.strip()

    # Centering Schemas: The paragraph containing "Рисунок"
    if text.startswith("Рисунок") or text.startswith("Схема"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0

    # Removing excess empty paragraphs
    if not text:
        empty_count += 1
        if empty_count >= 2: # Leave at most 1 empty line
            paragraphs_to_remove.append(p)
    else:
        empty_count = 0

for p in paragraphs_to_remove:
    p._element.getparent().remove(p._element)

# Remove trailing empty paragraphs at the end of the document to avoid empty pages
while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
    p = doc.paragraphs[-1]
    p._element.getparent().remove(p._element)

doc.save(output_path)
print("Finished updates.")
