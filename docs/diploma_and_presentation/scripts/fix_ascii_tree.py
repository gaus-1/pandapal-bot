import logging
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

DOC_PATH = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(DOC_PATH)

clean_tree = """frontend/src/
├─ features/         # Функциональные модули
│  ├─ AIChat/        # AI-чат со streaming
│  ├─ Achievements/  # Система достижений
│  ├─ Games/         # Образовательные игры
│  │  ├─ TicTacToe.tsx
│  │  ├─ Checkers.tsx
│  │  ├─ Game2048.tsx
│  │  └─ Erudite.tsx
│  ├─ MyPanda/       # Тамагочи - виртуальный питомец
│  ├─ Premium/       # Premium подписка
│  ├─ Donation/      # Система донатов
│  └─ Help/          # Помощь с уроками
├─ components/       # Переиспользуемые UI-компоненты
├─ hooks/            # Кастомные React-хуки
├─ services/         # API-клиенты и запросы
├─ store/            # Zustand - управление состоянием
├─ utils/            # Утилиты и константы
└─ types/            # TypeScript типы
"""

# Phase 1: Search and Destroy broken ascii tree
tree_start_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "frontend/src/" in p.text:
        tree_start_idx = i
        break

if tree_start_idx != -1:
    logger.info(f"Found tree starting at {tree_start_idx}")

    # Identify how many lines the broken tree spans
    end_idx = tree_start_idx
    for j in range(tree_start_idx + 1, tree_start_idx + 30):
        if j >= len(doc.paragraphs):
            break
        text = doc.paragraphs[j].text.strip()
        if "├" in text or "│" in text or "└" in text or text.startswith("features") or text.startswith("components") or "Zustand" in text or "TypeScript" in text:
            end_idx = j
        elif text == "":
            pass # Keep going over empty lines
        else:
            break

    logger.info(f"Tree spans from {tree_start_idx} to {end_idx}")

    anchor_elem = doc.paragraphs[tree_start_idx]._element

    # Phase 2: Insert clean block
    # Insert before the anchor, then delete the anchor and old lines
    for line in clean_tree.strip().split("\n"):
        new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        anchor_elem.addprevious(new_p_elem)
        paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(3.5)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(11)
        r_elem = run._r
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r_elem.insert(0, rPr)
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            'w:ascii="Courier New" w:hAnsi="Courier New" '
            'w:cs="Courier New" w:eastAsia="Courier New"/>'
        )
        rPr.append(rFonts)

    # Insert a space after the block
    space_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_elem.addprevious(space_p_elem)
    space_p = docx.text.paragraph.Paragraph(space_p_elem, doc._body)
    space_p.paragraph_format.space_after = Pt(12)

    # Delete old broken lines
    for j in range(tree_start_idx, end_idx + 1):
        p = doc.paragraphs[j]._element
        p.getparent().remove(p)

    doc.save(DOC_PATH)
    logger.info("Successfully fixed the ascii tree!")
else:
    logger.info("frontend/src/ not found.")
