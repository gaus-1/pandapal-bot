"""Verify the document after injection."""
import docx
from docx.oxml.ns import qn

doc = docx.Document(
    r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
)

# Count images
img_count = 0
for idx, p in enumerate(doc.paragraphs):
    ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    inlines = p._element.findall(f".//{{{ns}}}inline")
    if inlines:
        img_count += len(inlines)
        print(f"  Image at paragraph [{idx}]: alignment={p.alignment}")

print(f"\nTotal images: {img_count}")

print("\n=== FIGURE CAPTIONS ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Рисунок" in text and len(text) < 120 and ("—" in text or "А." in text):
        print(f"  [{i}] align={p.alignment} | {text}")

print("\n=== REMAINING ASCII BLOCKS ===")
box_count = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if any(c in text for c in "┌│└├─┐┘┤┬┴┼▼▲╔╗╚╝╠╣╦╩╬═║"):
        box_count += 1
print(f"  Remaining ASCII box-drawing paragraphs: {box_count}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
