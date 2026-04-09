import sys
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Загрузка документа
doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = Document(doc_path)

# Пути к сгенерированным картинкам
IMG_ERD = "docs/diploma_and_presentation/img/1_erd.png"
IMG_ARCH = "docs/diploma_and_presentation/img/2_architecture.png"
IMG_SEQ = "docs/diploma_and_presentation/img/3_sequence.png"
IMG_DEPLOY = "docs/diploma_and_presentation/img/4_deployment.png"
IMG_USECASE = "docs/diploma_and_presentation/img/5_usecase.png"

# 2. Форматирование текста и замена "10 ORM" на "15 ORM"
for p in doc.paragraphs:
    # Замена текста
    if "10 ORM-моделей" in p.text:
        p.text = p.text.replace("10 ORM-моделей", "15 ORM-моделей")
    if "10 ключевым моделям" in p.text:
        p.text = p.text.replace("10 ключевым моделям", "15 основным моделям")
    if "содержит 10 моделей" in p.text:
        p.text = p.text.replace("содержит 10 моделей", "содержит 15 моделей")

    # Применение ГОСТ-форматирования к обычным paragraphs
    if p.style.name == 'Normal':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Центрирование подписей рисунков/таблиц/приложений
    if p.text.strip().startswith("Рисунок") or p.text.strip().startswith("Таблица") or p.text.strip().startswith("Приложение") or p.text.strip().startswith("ГЛАВА"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        if p.runs:
           p.runs[0].font.name = 'Times New Roman'
           p.runs[0].font.size = Pt(14)


# 3. Вставка диаграмм вместо ASCII
# Мы пройдем по документу сверху вниз.
# Чтобы удалить ASCII блоки, мы найдем маркеры начала и конца или просто удалим все строки содержащие ┌, │, └, ─

def is_ascii_art(text):
    for char in ['┌', '├', '└', '│', '─', '┬', '┴', '┼', '◄', '►']:
        if char in text:
            return True
    return False

# Ищем маркеры для вставки
inserted_arch = False
inserted_deploy = False
inserted_erd = False
inserted_seq = False
inserted_usecase = False

for i, p in enumerate(doc.paragraphs):

    # Очистка ASCII art - просто очищаем текст в параграфах
    if is_ascii_art(p.text):
        p.clear()
        continue

    text = p.text.strip()

    if "архитектура системы может быть представлена в виде следующей диаграммы:" in text and not inserted_arch:
        run = p.add_run()
        run.add_picture(IMG_ARCH, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        inserted_arch = True

    if "Ниже представлена упрощённая ER-диаграмма" in text and not inserted_erd:
        run = p.add_run()
        run.add_picture(IMG_ERD, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        inserted_erd = True

    if "Приложение Б. Схема развёртывания платформы" in text and not inserted_deploy:
        run = p.add_run()
        run.add_picture(IMG_DEPLOY, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        inserted_deploy = True

# Добавим usecase и sequence во Введение / Главу 2, если найдем подходящее место
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Общий поток данных можно описать следующей схемой:" in text and not inserted_seq:
        run = p.add_run()
        run.add_picture(IMG_SEQ, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = 0
        inserted_seq = True

    if "Планируемый проектный продукт" in text and not inserted_usecase:
        # Вставим usecase перед этим
        # Insert a new paragraph with the image
        insert_p = p.insert_paragraph_before("Рисунок 1 — Диаграмма вариантов использования платформы (Use Case)")
        insert_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        img_p = p.insert_paragraph_before("")
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(IMG_USECASE, width=Inches(6.0))
        inserted_usecase = True

# Очистка пустых параграфов, которые остались от ASCII или пустых страниц
# В python-docx нельзя удалить параграф легко, но можно очистить текст.
# Но если слишком много пустых подряд, очистим.
empty_count = 0
for p in doc.paragraphs:
    if not p.text.strip():
        empty_count += 1
        if empty_count > 2:
            p.clear()
    else:
        empty_count = 0

# 4. Сохранение
output_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
doc.save(output_path)
print(f"Документ сохранен: {output_path}")
