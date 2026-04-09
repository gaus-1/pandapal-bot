"""
Comprehensive DOCX formatter for diploma project.
- Centers ALL diagrams (ASCII art) on the page
- Centers ALL figure captions (Рисунок X)
- Centers ALL table captions (Таблица X)
- Centers ALL tables
- Rebuilds ER diagram to be more beautiful and complete (14 tables)
- Applies proper ГОСТ formatting: monospace font, consistent sizing
- Ensures proper spacing and proportions
"""
import docx
import re
import copy
import logging
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
out_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"

doc = docx.Document(doc_path)

# ============================================================
# 1. CENTER ALL ASCII-ART DIAGRAM PARAGRAPHS
# ============================================================
logger.info("Step 1: Centering ASCII art diagrams...")

box_chars = set("┌│└├─┐┘┤┬┴┼▼▲►◄═║╔╗╚╝╠╣╦╩╬┘──────")
arrow_lines = {"↓", "↑", "→", "←"}

for p in doc.paragraphs:
    text = p.text.strip()

    # Check if this is ASCII art (box-drawing chars)
    has_box = any(c in text for c in "┌│└├─┐┘┤┬┴┼▼▲►◄═║╔╗╚╝╠╣╦╩╬")

    # Check if it's an arrow-only line (↓, →, etc.)
    is_arrow_line = text in arrow_lines or (
        len(text) <= 60 and
        text.replace(" ", "") in arrow_lines and
        not any(c.isalpha() and c not in "↓↑→←" for c in text)
    )

    # Check if it's a flow line like "Пользователь → Telegram / Браузер"
    is_flow_line = "→" in text or "↓" in text or "←" in text
    is_just_arrow = is_flow_line and len(text.replace(" ", "").replace("↓", "").replace("→", "").replace("←", "").replace("↑", "")) == 0

    if has_box or is_just_arrow:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)

        # Set monospace font for box-drawing
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            # Set East Asia font too
            r_elem = run._r
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                r_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New" w:eastAsia="Courier New"/>')
                rPr.insert(0, rFonts)
            else:
                rFonts.set(qn('w:ascii'), 'Courier New')
                rFonts.set(qn('w:hAnsi'), 'Courier New')
                rFonts.set(qn('w:cs'), 'Courier New')
                rFonts.set(qn('w:eastAsia'), 'Courier New')

        # Reduce spacing for compactness
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)

# ============================================================
# 2. CENTER FLOW/ARCHITECTURE TEXT LINES
# ============================================================
logger.info("Step 2: Centering flow/architecture text lines...")

# These are lines that are part of diagrams but contain text + arrows
flow_patterns = [
    "Пользователь →",
    "Telegram Bot API",
    "PostgreSQL",
    "Redis",
    "Yandex Cloud",
    "YooKassa",
    "+ pgvector",
    "(Upstash)",
    "(GPT, STT,",
    "OCR, ART)",
    "(платежи)",
]

for p in doc.paragraphs:
    text = p.text.strip()

    # Lines that are part of flow diagrams
    if text and any(pattern in text for pattern in flow_patterns):
        # Don't touch regular body text paragraphs
        if p.alignment not in (WD_ALIGN_PARAGRAPH.JUSTIFY,) or len(text) < 80:
            if not text.startswith(("•", "1.", "2.", "3.", "4.", "5.", "6.")) and len(text) < 80:
                # Check if it really looks like a diagram line (short, no periods at end)
                if not text.endswith(".") and not text.startswith("В ") and not text.startswith("Платформа"):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Pt(0)
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = Pt(10)

# ============================================================
# 3. CENTER ALL FIGURE CAPTIONS
# ============================================================
logger.info("Step 3: Centering figure captions...")

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("Рисунок") and "—" in text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.italic = True

# ============================================================
# 4. CENTER TABLE CAPTIONS
# ============================================================
logger.info("Step 4: Centering table captions...")

for p in doc.paragraphs:
    text = p.text.strip()
    if re.match(r'^Таблица\s+(А\.)?\d+\s*[—–-]', text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

# ============================================================
# 5. CENTER ALL TABLES
# ============================================================
logger.info("Step 5: Centering all tables...")

for table in doc.tables:
    # Set table alignment to center
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    # Remove existing jc
    for jc in tblPr.findall(qn('w:jc')):
        tblPr.remove(jc)

    # Add center alignment
    jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>')
    tblPr.append(jc)

    # Style each cell for beauty
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    # Style header row (first row) differently
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            # Add shading to header
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)

            # Add light blue shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="D9E2F3"/>')
            # Remove existing shading
            for old_shd in tcPr.findall(qn('w:shd')):
                tcPr.remove(old_shd)
            tcPr.append(shading)

            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

# ============================================================
# 6. FIX CODE BLOCKS - proper monospace formatting
# ============================================================
logger.info("Step 6: Formatting code blocks...")

code_markers = [
    "class User(Base):",
    "__tablename__",
    "Mapped[",
    "mapped_column(",
    "def build_system_prompt",
    "def moderate_message",
    "async def ",
    "def is_premium",
    "@property",
    "return ModerationResult",
    "return f\"\"\"",
    "forbidden_patterns_check",
    "advanced_moderation_check",
    "ai_moderator.check",
    "adult_topics_service",
    "emoji_policy",
    "ЗАПРЕЩЕНО:",
    "style = \"",
]

for p in doc.paragraphs:
    text = p.text.strip()
    if any(marker in text for marker in code_markers):
        # This is a code line
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# ============================================================
# 7. FIX FRONTEND STRUCTURE TREE
# ============================================================
logger.info("Step 7: Formatting frontend structure tree...")

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith(("├──", "│", "└──", "frontend/src/")) or (
        "├──" in text or "└──" in text or ("│" in text and "#" in text)
    ):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# ============================================================
# 8. FIX PROCESS FLOW (text-based flow diagrams)
# ============================================================
logger.info("Step 8: Formatting text-based flow diagrams...")

# The text-based flow like:
# "bot/handlers/ai_chat/text.py → получение текстового сообщения"
for p in doc.paragraphs:
    text = p.text.strip()
    if "→" in text and ("bot/" in text or "Пользователь" in text):
        if len(text) < 100 and not text.endswith("."):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    # Arrow-only lines
    if text == "↓" or text == "→":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# ============================================================
# 9. REBUILD ER DIAGRAM - Beautiful version with ALL 14 tables
# ============================================================
logger.info("Step 9: Rebuilding ER diagram...")

# Find the ER diagram in Приложение А
er_start = -1
er_end = -1
found_appendix = False

for i, p in enumerate(doc.paragraphs):
    if "База данных PandaPal содержит 14 ORM-моделей" in p.text:
        found_appendix = True
        continue

    if found_appendix and any(c in p.text for c in "┌│└├─┐┘▼▲"):
        if er_start == -1:
            er_start = i
        er_end = i
    elif found_appendix and er_start != -1 and not any(c in p.text for c in "┌│└├─┐┘▼▲"):
        # Check if next has box chars (in case of an empty line between diagram parts)
        if p.text.strip() == "" and i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i + 1].text
            if any(c in next_text for c in "┌│└├─┐┘▼▲"):
                er_end = i  # include empty line
                continue
        break

# Check if we also have a Рисунок caption right after
if er_end != -1 and er_end + 1 < len(doc.paragraphs):
    next_p = doc.paragraphs[er_end + 1]
    if next_p.text.strip().startswith("Рисунок 3"):
        er_end = er_end + 1  # Include the caption in deletion

logger.info(f"ER diagram found: paragraphs {er_start} to {er_end}")

if er_start != -1 and er_end != -1:
    # Delete old ER diagram
    count = er_end - er_start + 1
    for _ in range(count):
        p = doc.paragraphs[er_start]
        p._element.getparent().remove(p._element)

    # Insert new beautiful ER diagram
    # The design: users at center top, 5 direct children, each with their own children
    new_er = [
        "╔══════════════════════════════════════════════════════════════════════════════════╗",
        "║                            СХЕМА БАЗЫ ДАННЫХ (ERD)                             ║",
        "║                     14 ORM-моделей, связанных внешними ключами                  ║",
        "╚══════════════════════════════════════════════════════════════════════════════════╝",
        "",
        "                           ┌────────────────────┐",
        "                           │       USERS        │",
        "                           │────────────────────│",
        "                           │ id (PK)            │",
        "                           │ telegram_id (UQ)   │",
        "                           │ username           │",
        "                           │ first_name         │",
        "                           │ age, grade         │",
        "                           │ user_type          │",
        "                           │ is_premium         │",
        "                           │ total_coins        │",
        "                           │ created_at         │",
        "                           └─────────┬──────────┘",
        "             ┌──────────┬────────────┼────────────┬──────────┐",
        "             │          │            │            │          │",
        "             ▼          ▼            ▼            ▼          ▼",
        "    ┌──────────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐",
        "    │  PANDA_PETS  │ │REFERRALS │ │CHAT_HIST │ │ LEARNING │ │  GAMES   │",
        "    │──────────────│ │──────────│ │──────────│ │──────────│ │──────────│",
        "    │id (PK)       │ │id (PK)   │ │id (PK)   │ │id (PK)   │ │id (PK)   │",
        "    │user_id (FK)──│ │user_id(FK│ │user_id(FK│ │user_id(FK│ │user_id(FK│",
        "    │name          │ │ref_code  │ │role      │ │subject   │ │game_type │",
        "    │health (0-100)│ │bonus     │ │content   │ │progress  │ │score     │",
        "    │energy (0-100)│ │status    │ │tokens    │ │xp_earned │ │result    │",
        "    │satiety       │ │created_at│ │created_at│ │completed │ │played_at │",
        "    │mood (0-100)  │ └──────────┘ └──────────┘ └──────────┘ └──────────┘",
        "    │last_fed_at   │",
        "    └──────┬───────┘",
        "           │",
        "     ┌─────┴─────┐           ┌─────────────────────────────────────────┐",
        "     │            │           │  СВЯЗИ МЕЖДУ СУЩНОСТЯМИ:               │",
        "     ▼            ▼           │                                         │",
        " ┌──────────┐ ┌──────────┐   │  users ─┬─► panda_pets    (1:1)         │",
        " │ PAYMENTS │ │ PAYOUTS  │   │         ├─► referrals     (1:N)         │",
        " │──────────│ │──────────│   │         ├─► chat_history  (1:N)         │",
        " │id (PK)   │ │id (PK)   │   │         ├─► learning      (1:N)         │",
        " │user_id(FK│ │user_id(FK│   │         ├─► games         (1:N)         │",
        " │amount    │ │amount    │   │         ├─► payments      (1:N)         │",
        " │status    │ │status    │   │         ├─► payouts       (1:N)         │",
        " │provider  │ │method    │   │         ├─► daily_bonuses (1:N)         │",
        " │created_at│ │created_at│   │         ├─► homework      (1:N)         │",
        " └──────────┘ └──────────┘   │         ├─► game_stats    (1:N)         │",
        "                             │         ├─► subscriptions (1:N)         │",
        "     ┌──────────────────┐    │         └─► analytics     (1:N)         │",
        "     │                  │    │                                         │",
        "     ▼                  ▼    │  panda_pets ──► payments   (indirect)   │",
        " ┌──────────┐ ┌──────────┐   │  panda_pets ──► payouts    (indirect)   │",
        " │  DAILY   │ │ HOMEWORK │   └─────────────────────────────────────────┘",
        " │──────────│ │──────────│",
        " │id (PK)   │ │id (PK)   │",
        " │user_id(FK│ │user_id(FK│         ┌──────────────────────────────┐",
        " │bonus_amt │ │task_text │         │  ДОПОЛНИТЕЛЬНЫЕ СУЩНОСТИ     │",
        " │is_claimed│ │subject   │         │──────────────────────────────│",
        " │date      │ │photo_url │         │  ┌────────────┐             │",
        " └──────────┘ │ai_review │         │  │ GAME_STATS │             │",
        "              │score     │         │  │────────────│             │",
        "              │created_at│         │  │id (PK)     │             │",
        "              └──────────┘         │  │user_id (FK)│             │",
        "                                  │  │game_type   │             │",
        " ┌──────────┐ ┌──────────┐        │  │wins, draws │             │",
        " │   SUBS   │ │ANALYTICS │        │  │losses      │             │",
        " │──────────│ │──────────│        │  │best_score  │             │",
        " │id (PK)   │ │id (PK)   │        │  └────────────┘             │",
        " │user_id(FK│ │user_id(FK│        │                             │",
        " │plan_type │ │event_type│        │  ┌────────────────────┐     │",
        " │starts_at │ │metadata  │        │  │KNOWLEDGE_EMBEDDINGS│     │",
        " │expires_at│ │created_at│        │  │────────────────────│     │",
        " │is_active │ └──────────┘        │  │id (PK)            │     │",
        " │auto_renew│                     │  │content             │     │",
        " └──────────┘                     │  │subject             │     │",
        "                                  │  │embedding (vector)  │     │",
        "                                  │  │created_at          │     │",
        "                                  │  └────────────────────┘     │",
        "                                  └──────────────────────────────┘",
    ]

    # Insert the diagram
    anchor = doc.paragraphs[er_start]

    for line in new_er:
        new_p = docx.oxml.OxmlElement('w:p')
        anchor._element.addprevious(new_p)
        para = docx.text.paragraph.Paragraph(new_p, doc._body)

        if line == "":
            para.text = ""
            continue

        para.text = line
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(9)

        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(7)
            # Ensure monospace rendering
            r_elem = run._r
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                r_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New" w:eastAsia="Courier New"/>')
                rPr.insert(0, rFonts)
            else:
                rFonts.set(qn('w:ascii'), 'Courier New')
                rFonts.set(qn('w:hAnsi'), 'Courier New')
                rFonts.set(qn('w:cs'), 'Courier New')
                rFonts.set(qn('w:eastAsia'), 'Courier New')

    # Add figure caption after diagram
    cap_p = docx.oxml.OxmlElement('w:p')
    anchor._element.addprevious(cap_p)
    caption = docx.text.paragraph.Paragraph(cap_p, doc._body)
    caption.text = "Рисунок 3 — ER-диаграмма базы данных PandaPal (14 сущностей)"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(12)
    for run in caption.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.italic = True

# ============================================================
# 10. CENTER the "Приложение" headings & appendix titles
# ============================================================
logger.info("Step 10: Centering appendix headings...")

for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("Приложение") and len(text) < 100:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True

# ============================================================
# 11. CENTER "ПРИЛОЖЕНИЕ А", "(обязательное)", "Схема базы данных (ERD)"
# ============================================================
logger.info("Step 11: Centering supplementary heading blocks...")

for p in doc.paragraphs:
    text = p.text.strip()
    if text in ("ПРИЛОЖЕНИЕ А", "ПРИЛОЖЕНИЕ Б", "(обязательное)",
                "Схема базы данных (ERD)", "Архитектура программного комплекса"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

# ============================================================
# 12. FIX Appendix Б architecture diagram - make it centered & beautiful
# ============================================================
logger.info("Step 12: Re-formatting Appendix Б architecture text...")

found_appendix_b = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "ПРИЛОЖЕНИЕ Б" in text:
        found_appendix_b = True
        continue

    if found_appendix_b and text.startswith("Архитектура системы построена"):
        # This is a long text paragraph - center it or make it JUSTIFY
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

# ============================================================
# 13. ENSURE SPACING CONSISTENCY for body text
# ============================================================
logger.info("Step 13: Ensuring spacing consistency...")

for p in doc.paragraphs:
    text = p.text.strip()

    # Skip empty paragraphs
    if not text:
        continue

    # Body text should have 1.5 line spacing per ГОСТ
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        if not any(c in text for c in "┌│└├─┐┘") and "→" not in text:
            # Regular body text
            p.paragraph_format.line_spacing = Pt(21)  # 1.5 spacing for 14pt

    # Section headings (ГЛАВА, ВВЕДЕНИЕ, etc.)
    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        if any(h in text for h in ["ГЛАВА", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СОДЕРЖАНИЕ",
                                    "СПИСОК ИСПОЛЬЗОВАННЫХ", "ПРИЛОЖЕНИЯ", "ПАСПОРТ ПРОЕКТА",
                                    "СПИСОК СОКРАЩЕНИЙ", "ДИПЛОМНАЯ РАБОТА"]):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                run.font.bold = True

# ============================================================
# 14. ADD TABLE BORDERS (clean ГОСТ-style single borders)
# ============================================================
logger.info("Step 14: Adding clean table borders...")

def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with dict of sz, val, color."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

    for edge, attrs in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" w:sz="{attrs.get("sz", "4")}" w:color="{attrs.get("color", "000000")}" w:space="0"/>')
            tcBorders.append(element)
        else:
            element.set(qn('w:val'), attrs.get('val', 'single'))
            element.set(qn('w:sz'), attrs.get('sz', '4'))
            element.set(qn('w:color'), attrs.get('color', '000000'))

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"sz": "4", "val": "single", "color": "000000"},
                bottom={"sz": "4", "val": "single", "color": "000000"},
                left={"sz": "4", "val": "single", "color": "000000"},
                right={"sz": "4", "val": "single", "color": "000000"},
            )

# ============================================================
# SAVE
# ============================================================
logger.info("Saving document...")
doc.save(out_path)
logger.info(f"Document saved to {out_path}")
