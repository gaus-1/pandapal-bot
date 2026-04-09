"""
Inject professional PNG diagrams into the diploma DOCX.

Strategy:
 - Work from END to START of the document (reverse order) to avoid index shifts.
 - For each diagram section: delete old ASCII paragraphs, then insert a
   centered image + updated caption at the same position.
 - Add NEW diagrams (use case, sequence) that don't exist yet.

Diagrams to inject:
  1. Figure 1 — Architecture (replaces old biz-process diagram)  → 2_architecture.png
  2. Figure 2 — Use Case diagram (NEW)                          → 5_usecase.png
  3. Figure 3 — Sequence diagram (NEW)                          → 3_sequence.png
  4. Figure 4 — Deployment diagram                               → 4_deployment.png
  5. Appendix A — ERD conceptual + physical                      → 1_erd.png + 6_erd_physical.png
  6. Appendix B — Deployment                                      → 4_deployment.png
  7. Appendix (relational schema)                                 → 7_erd_relational.png
"""
import logging
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

DOC_PATH = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
IMG_DIR = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\img"

doc = docx.Document(DOC_PATH)


# ── HELPERS ──────────────────────────────────────────────────────────────────

def is_box_char_line(text: str) -> bool:
    """Check if a string contains ASCII box-drawing characters."""
    return any(c in text for c in "┌│└├─┐┘┤┬┴┼▼▲╔╗╚╝╠╣╦╩╬═║")


def delete_paragraph(paragraph):
    """Delete a paragraph from the document."""
    p = paragraph._element
    p.getparent().remove(p)


def insert_image_paragraph(doc, anchor_element, img_path, width_cm=15):
    """Insert a centered image paragraph BEFORE the anchor element."""
    new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_element.addprevious(new_p_elem)
    paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return paragraph


def insert_caption_paragraph(doc, anchor_element, caption_text):
    """Insert a centered italic caption paragraph BEFORE the anchor element."""
    new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_element.addprevious(new_p_elem)
    paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(caption_text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.italic = False
    # Set font for cyrillic
    r_elem = run._r
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} '
                'w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
                'w:cs="Times New Roman" w:eastAsia="Times New Roman"/>'
            )
            rPr.insert(0, rFonts)
    return paragraph


def insert_text_paragraph(doc, anchor_element, text, font_size=14, bold=False, align="center"):
    """Insert a text paragraph BEFORE the anchor element."""
    new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_element.addprevious(new_p_elem)
    paragraph = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(0) if align == "center" else Cm(1.25)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return paragraph


# ── PHASE 1: IDENTIFY ALL DIAGRAM BLOCKS ────────────────────────────────────

logger.info("Phase 1: Identifying diagram blocks...")

# Collect info
paragraphs = doc.paragraphs

# Find all contiguous blocks of ASCII box-drawing / arrow / flow paragraphs
# that belong to each diagram section.

# Diagram 1: "Рисунок 1" area — paragraphs from ~307 to ~335
# We need to find the exact boundaries dynamically.

def find_caption_idx(caption_start: str) -> int:
    """Find paragraph index of a caption starting with given text."""
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith(caption_start):
            return i
    return -1


def find_diagram_block_before(caption_idx: int) -> tuple:
    """Find the start of the diagram block (ASCII + arrow lines) before a caption."""
    start = caption_idx - 1
    while start >= 0:
        text = paragraphs[start].text.strip()
        has_box = is_box_char_line(text)
        is_arrow = text in ("↓", "→", "↑", "←")
        is_flow = "→" in text and len(text) < 80
        is_empty = text == ""
        is_short_label = len(text) < 60 and not text.endswith(".") and not text.startswith(("•", "ГЛАВА", "Таблица"))

        if has_box or is_arrow or is_flow or is_empty:
            start -= 1
        elif is_short_label and start > caption_idx - 30:
            # Could be a label like "PostgreSQL  Redis" inside the diagram
            start -= 1
        else:
            break
    return start + 1


def find_diagram_block_after(caption_idx: int) -> int:
    """Find the end of remaining diagram elements after a caption."""
    end = caption_idx + 1
    while end < len(paragraphs):
        text = paragraphs[end].text.strip()
        has_box = is_box_char_line(text)
        is_arrow = text in ("↓", "→", "↑", "←")
        is_flow = "→" in text and len(text) < 80 and not text.endswith(".")
        is_empty = text == ""
        is_short_label = len(text) < 60 and not text.endswith(".") and not text.startswith(("•", "ГЛАВА", "Таблица", "Приложение"))

        if has_box or is_arrow:
            end += 1
        elif is_empty and end < caption_idx + 5:
            end += 1
        elif is_short_label and is_flow:
            end += 1
        else:
            break
    return end - 1


# ── IDENTIFY DIAGRAM SECTIONS ───────────────────────────────────────────────

# Section 1: Figure 1 + continuation below it
fig1_caption = find_caption_idx("Рисунок 1")
fig1_block_start = find_diagram_block_before(fig1_caption)
fig1_block_end_after = find_diagram_block_after(fig1_caption)

# Section 2: Figure 2 + continuation below it
fig2_caption = find_caption_idx("Рисунок 2")
fig2_block_start = find_diagram_block_before(fig2_caption)
fig2_block_end_after = find_diagram_block_after(fig2_caption)

# Section 3: ER diagram in Appendix A
fig3_caption = find_caption_idx("Рисунок 3")
# Find the title line for Appendix A
appendix_a_idx = -1
for i, p in enumerate(paragraphs):
    if "Приложение А" in p.text.strip() and "Схема" in p.text.strip():
        appendix_a_idx = i
        break

# Section 4: Deployment in Appendix B
fig4_caption = find_caption_idx("Рисунок 4")
appendix_b_idx = -1
for i, p in enumerate(paragraphs):
    if "Приложение Б" in p.text.strip() and "Схема" in p.text.strip():
        appendix_b_idx = i
        break

logger.info(f"  Figure 1: caption={fig1_caption}, block={fig1_block_start}-{fig1_block_end_after}")
logger.info(f"  Figure 2: caption={fig2_caption}, block={fig2_block_start}-{fig2_block_end_after}")
logger.info(f"  Figure 3 (ERD): caption={fig3_caption}, appendix_a={appendix_a_idx}")
logger.info(f"  Figure 4 (Deploy): caption={fig4_caption}, appendix_b={appendix_b_idx}")


# ── PHASE 2: REPLACE DIAGRAMS (work from end to start) ──────────────────────

logger.info("\nPhase 2: Replacing diagram blocks with images...")

# Helper: delete a range of paragraphs
def delete_range(start_idx: int, end_idx: int):
    """Delete paragraphs from start_idx to end_idx inclusive."""
    count = 0
    for idx in range(end_idx, start_idx - 1, -1):
        if idx < len(doc.paragraphs):
            delete_paragraph(doc.paragraphs[idx])
            count += 1
    return count


# ── 4. APPENDIX B: Deployment diagram ───────────────────────────────────────
logger.info("\n--- Replacing Appendix B: Deployment diagram ---")

# Delete old ASCII diagram + old caption from Appendix B
# Block is from line after "Приложение Б" heading to the fig4_caption
if fig4_caption > 0 and appendix_b_idx > 0:
    # Delete from appendix_b_idx+1 (first box line) through fig4_caption (old caption)
    del_start = appendix_b_idx + 1
    del_end = fig4_caption
    deleted = delete_range(del_start, del_end)
    logger.info(f"  Deleted {deleted} paragraphs [{del_start}-{del_end}]")

    # Now insert image + caption at the position where the deleted block was
    # The anchor is the paragraph currently at del_start (which is what was after the deleted range)
    anchor = doc.paragraphs[del_start] if del_start < len(doc.paragraphs) else doc.paragraphs[-1]
    anchor_elem = anchor._element

    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок 4 – Схема развёртывания инфраструктуры PandaPal")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "4_deployment.png"), width_cm=16)
    logger.info("  Inserted deployment image + caption")


# ── 3. APPENDIX A: ERD diagram ─────────────────────────────────────────────
logger.info("\n--- Replacing Appendix A: ERD diagram ---")

# Re-find indices after deletions
fig3_caption = find_caption_idx("Рисунок 3")
appendix_a_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Приложение А" in p.text.strip() and "Схема" in p.text.strip():
        appendix_a_idx = i
        break

if fig3_caption > 0 and appendix_a_idx > 0:
    # Find the text line "База данных PandaPal содержит..."
    desc_idx = appendix_a_idx + 1

    # Delete from desc_idx+1 (first box) through fig3_caption (old caption)
    # But first find the actual first box line
    first_box = desc_idx + 1
    while first_box < fig3_caption:
        text = doc.paragraphs[first_box].text.strip()
        if is_box_char_line(text) or text == "":
            break
        first_box += 1

    del_start = first_box
    del_end = fig3_caption
    deleted = delete_range(del_start, del_end)
    logger.info(f"  Deleted {deleted} paragraphs [{del_start}-{del_end}]")

    # Update the description text
    desc_p = doc.paragraphs[desc_idx]
    desc_p.clear()
    run = desc_p.add_run(
        "База данных PandaPal содержит 15 ORM-моделей, связанных внешними ключами. "
        "Ниже представлена концептуальная ER-диаграмма, отображающая основные сущности "
        "и связи между ними, а также физическая модель с перечнем всех атрибутов."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    desc_p.paragraph_format.first_line_indent = Cm(1.25)

    # Insert images: conceptual ERD + physical ERD + relational schema
    anchor = doc.paragraphs[del_start] if del_start < len(doc.paragraphs) else doc.paragraphs[-1]
    anchor_elem = anchor._element

    # Insert in reverse order (BEFORE anchor each time, so first inserted = last in doc)

    # 3c. Relational schema
    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок А.3 – Реляционная схема базы данных PandaPal")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "7_erd_relational.png"), width_cm=17)

    # 3b. Physical ERD
    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок А.2 – Физическая ER-диаграмма базы данных (15 таблиц)")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "6_erd_physical.png"), width_cm=17)

    # 3a. Conceptual ERD
    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок А.1 – Концептуальная ER-диаграмма базы данных PandaPal")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "1_erd.png"), width_cm=16)

    logger.info("  Inserted 3 ERD images + captions")


# ── 2. FIGURE 2: Architecture + data flow ──────────────────────────────────
logger.info("\n--- Replacing in-text Figure 2: Architecture ---")

# Re-find
fig2_caption = find_caption_idx("Рисунок 2")
if fig2_caption > 0:
    fig2_start = find_diagram_block_before(fig2_caption)
    fig2_end = find_diagram_block_after(fig2_caption)

    del_start = fig2_start
    del_end = fig2_end
    deleted = delete_range(del_start, del_end)
    logger.info(f"  Deleted {deleted} paragraphs [{del_start}-{del_end}]")

    anchor = doc.paragraphs[del_start] if del_start < len(doc.paragraphs) else doc.paragraphs[-1]
    anchor_elem = anchor._element

    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок 2 – Общая архитектура системы PandaPal")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "2_architecture.png"), width_cm=16)
    logger.info("  Inserted architecture image + caption")


# ── 1. FIGURE 1: Business processes ────────────────────────────────────────
logger.info("\n--- Replacing in-text Figure 1: Business processes ---")

# Re-find
fig1_caption = find_caption_idx("Рисунок 1")
if fig1_caption > 0:
    fig1_start = find_diagram_block_before(fig1_caption)
    fig1_end = find_diagram_block_after(fig1_caption)

    del_start = fig1_start
    del_end = fig1_end
    deleted = delete_range(del_start, del_end)
    logger.info(f"  Deleted {deleted} paragraphs [{del_start}-{del_end}]")

    anchor = doc.paragraphs[del_start] if del_start < len(doc.paragraphs) else doc.paragraphs[-1]
    anchor_elem = anchor._element

    # Insert: Use Case + Sequence + Architecture overview for main body
    # Actually for Figure 1 we keep it as diagram of use-case
    insert_caption_paragraph(doc, anchor_elem,
                             "Рисунок 1 – Диаграмма вариантов использования PandaPal")
    insert_image_paragraph(doc, anchor_elem,
                           os.path.join(IMG_DIR, "5_usecase.png"), width_cm=14)
    logger.info("  Inserted use-case image + caption")


# ── PHASE 3: ADD SEQUENCE DIAGRAM after Figure 2 ────────────────────────────
logger.info("\n--- Adding sequence diagram as Figure 3 ---")

# Find architecture caption we just inserted
fig2_new = find_caption_idx("Рисунок 2")
if fig2_new > 0:
    # Insert sequence diagram right after Figure 2 caption
    anchor_idx = fig2_new + 1
    if anchor_idx < len(doc.paragraphs):
        anchor = doc.paragraphs[anchor_idx]
        anchor_elem = anchor._element

        insert_caption_paragraph(doc, anchor_elem,
                                 "Рисунок 3 – Диаграмма последовательности обработки сообщений")
        insert_image_paragraph(doc, anchor_elem,
                               os.path.join(IMG_DIR, "3_sequence.png"), width_cm=15)
        logger.info("  Inserted sequence diagram as Figure 3")


# ── PHASE 4: FIX REMAINING TEXT REFERENCES ──────────────────────────────────
logger.info("\nPhase 4: Final cleanup...")

# Update old references: "14 ORM-моделей" -> "15 ORM-моделей"
for p in doc.paragraphs:
    if "14 ORM-модел" in p.text or "14 сущност" in p.text:
        for run in p.runs:
            if "14 ORM-модел" in run.text:
                run.text = run.text.replace("14 ORM-модел", "15 ORM-модел")
                logger.info(f"  Fixed: 14 -> 15 ORM models")
            if "14 сущност" in run.text:
                run.text = run.text.replace("14 сущност", "15 сущност")
                logger.info(f"  Fixed: 14 -> 15 entities")

# Make sure all inline images are centered
for p in doc.paragraphs:
    # Check if paragraph contains an image
    if p._element.findall(f'.//{qn("wp:inline")}') or p._element.findall(f'.//{qn("wp:anchor")}'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)


# ── SAVE ────────────────────────────────────────────────────────────────────
logger.info("\nSaving document...")
doc.save(DOC_PATH)
logger.info(f"Document saved to {DOC_PATH}")
logger.info("Done! All 7 professional diagrams injected.")
