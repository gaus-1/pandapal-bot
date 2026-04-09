"""
Precise surgical fix for diagram ordering and cleanup.
Works with XML elements directly to rearrange paragraphs.
"""
import logging
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

DOC_PATH = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
IMG_DIR = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\img"

doc = docx.Document(DOC_PATH)

WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def has_image(p):
    return bool(p._element.findall(f".//{{{WP_NS}}}inline"))


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def find_idx(text_contains, after_idx=0):
    """Find paragraph index containing text, searching after a given index."""
    for i, p in enumerate(doc.paragraphs):
        if i >= after_idx and text_contains in p.text:
            return i
    return -1


def insert_image_after(anchor_p, img_path, width_cm=15):
    """Insert centered image AFTER anchor paragraph. Returns the new paragraph."""
    new_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_p._element.addnext(new_elem)
    para = docx.text.paragraph.Paragraph(new_elem, doc._body)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return para


def insert_caption_after(anchor_p, text):
    """Insert centered italic caption AFTER anchor paragraph. Returns the new paragraph."""
    new_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    anchor_p._element.addnext(new_elem)
    para = docx.text.paragraph.Paragraph(new_elem, doc._body)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.italic = True
    rPr = run._r.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            'w:ascii="Times New Roman" w:hAnsi="Times New Roman" '
            'w:cs="Times New Roman" w:eastAsia="Times New Roman"/>'
        )
        rPr.insert(0, rFonts)
    return para


# ════════════════════════════════════════════════════════════════════════════
# FIX 1: Clean up leftover text after Figure 1
# ════════════════════════════════════════════════════════════════════════════
logger.info("FIX 1: Cleaning up leftover text after Figure 1...")

# Find leftover lines: "PostgreSQL  Redis...", "+ pgvector...", "OCR, ART)"
# They should be right after the Figure 1 image
fig1_cap_idx = find_idx("Рисунок 1")
if fig1_cap_idx > 0:
    # Check paragraphs after image for leftover text
    check_start = fig1_cap_idx + 2  # skip caption and image
    deleted_count = 0
    while check_start < len(doc.paragraphs):
        text = doc.paragraphs[check_start].text.strip()
        # These are leftover diagram labels
        if text in (
            "PostgreSQL  Redis      Yandex Cloud  YooKassa",
            "+ pgvector  (Upstash)  (GPT, STT,   (платежи)",
            "+ pgvector  (Upstash)  (GPT, STT,\xa0\xa0 (платежи)",
            "OCR, ART)",
        ) or (text.startswith("PostgreSQL") and "Redis" in text) or (
            text.startswith("+ pgvector")
        ) or text == "OCR, ART)":
            delete_paragraph(doc.paragraphs[check_start])
            deleted_count += 1
        else:
            break
    logger.info(f"  Deleted {deleted_count} leftover lines after Figure 1")


# ════════════════════════════════════════════════════════════════════════════
# FIX 2: Fix Figure 2 + Figure 3 area (captions and images interleaved)
# ════════════════════════════════════════════════════════════════════════════
logger.info("\nFIX 2: Fixing Figure 2 and Figure 3 area...")

# Current state: [502] caption2, [503] caption3, [504] image, [505] image
# Target:  [X] image2, [X+1] caption2, [X+2] image3, [X+3] caption3

# Find the intro text "Общая архитектура системы может быть представлена"
intro_idx = find_idx("может быть представлена в виде следующей диаграммы")
if intro_idx < 0:
    intro_idx = find_idx("может быть представлена в виде")

fig2_cap_idx = find_idx("Рисунок 2")
fig3_cap_idx = find_idx("Рисунок 3")

logger.info(f"  intro={intro_idx}, fig2_cap={fig2_cap_idx}, fig3_cap={fig3_cap_idx}")

if fig2_cap_idx > 0 and fig3_cap_idx > 0:
    # Delete old captions and images (4 paragraphs: cap2, cap3, img, img)
    # Find the end: first paragraph after images that isn't an image/caption
    del_start = fig2_cap_idx
    del_end = fig3_cap_idx

    # Count images right after the captions
    check = fig3_cap_idx + 1
    while check < len(doc.paragraphs) and has_image(doc.paragraphs[check]):
        del_end = check
        check += 1

    # Delete from del_start to del_end inclusive
    count = del_end - del_start + 1
    logger.info(f"  Deleting paragraphs [{del_start}-{del_end}] ({count} items)")
    for _ in range(count):
        delete_paragraph(doc.paragraphs[del_start])

    # Now insert properly: image, caption, image, caption
    # The anchor is the paragraph at del_start (which was the one after the deleted range)
    # Actually, we should insert after the intro text
    intro_idx = find_idx("может быть представлена в виде")
    if intro_idx < 0:
        intro_idx = find_idx("3.4. Визуализация")
        intro_idx += 1  # text after heading

    anchor = doc.paragraphs[intro_idx]

    # Insert in order using addnext chain:
    # 1. Architecture image
    p1 = insert_image_after(anchor, os.path.join(IMG_DIR, "2_architecture.png"), width_cm=16)
    # 2. Architecture caption
    p2 = insert_caption_after(p1, "Рисунок 2 — Общая архитектура системы PandaPal")
    # 3. Sequence image
    p3 = insert_image_after(p2, os.path.join(IMG_DIR, "3_sequence.png"), width_cm=15)
    # 4. Sequence caption
    p4 = insert_caption_after(p3, "Рисунок 3 — Диаграмма последовательности обработки сообщений")

    logger.info("  Inserted Figure 2 + Figure 3 correctly")


# ════════════════════════════════════════════════════════════════════════════
# FIX 3: Fix Appendix A ERD order (currently A.3, A.2, A.1 — need A.1, A.2, A.3)
# ════════════════════════════════════════════════════════════════════════════
logger.info("\nFIX 3: Fixing Appendix A ERD order...")

# Find Appendix A description (use unique text)
appendix_desc_idx = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if ("15 ORM-модел" in text and "концепту" in text and
            i > 400):  # Must be in appendix area, not TOC
        appendix_desc_idx = i
        break

# Find Appendix B heading (boundary)
appendix_b_idx = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("Приложение Б") and "Схема" in text and i > 400:
        appendix_b_idx = i
        break

logger.info(f"  Appendix A desc: [{appendix_desc_idx}], Appendix B: [{appendix_b_idx}]")

if appendix_desc_idx > 0 and appendix_b_idx > 0:
    # Delete everything between desc and Appendix B
    del_start = appendix_desc_idx + 1
    del_count = appendix_b_idx - del_start

    logger.info(f"  Deleting {del_count} old ERD paragraphs [{del_start}-{appendix_b_idx - 1}]")
    for _ in range(del_count):
        delete_paragraph(doc.paragraphs[del_start])

    # Re-insert in correct order using addnext
    anchor = doc.paragraphs[appendix_desc_idx]

    # A.1 — Conceptual
    p1 = insert_image_after(anchor, os.path.join(IMG_DIR, "1_erd.png"), width_cm=16)
    p2 = insert_caption_after(p1, "Рисунок А.1 — Концептуальная ER-диаграмма базы данных PandaPal")

    # A.2 — Physical
    p3 = insert_image_after(p2, os.path.join(IMG_DIR, "6_erd_physical.png"), width_cm=17)
    p4 = insert_caption_after(p3, "Рисунок А.2 — Физическая ER-диаграмма базы данных (15 таблиц)")

    # A.3 — Relational
    p5 = insert_image_after(p4, os.path.join(IMG_DIR, "7_erd_relational.png"), width_cm=17)
    p6 = insert_caption_after(p5, "Рисунок А.3 — Реляционная схема базы данных PandaPal")

    logger.info("  Inserted ERDs in correct order: A.1, A.2, A.3")


# ════════════════════════════════════════════════════════════════════════════
# FIX 4: Fix Figure 1 — swap caption and image (image should be first)
# ════════════════════════════════════════════════════════════════════════════
logger.info("\nFIX 4: Fixing Figure 1 order...")

fig1_new_idx = find_idx("Рисунок 1")
if fig1_new_idx > 0:
    cap_p = doc.paragraphs[fig1_new_idx]
    # Check if next is image
    if fig1_new_idx + 1 < len(doc.paragraphs) and has_image(doc.paragraphs[fig1_new_idx + 1]):
        img_p = doc.paragraphs[fig1_new_idx + 1]
        # Swap: move caption after image
        img_p._element.addnext(cap_p._element)
        logger.info("  Swapped Figure 1: image now above caption")
    else:
        logger.info("  Figure 1 already in correct order")


# ════════════════════════════════════════════════════════════════════════════
# FINAL: Ensure all images centered
# ════════════════════════════════════════════════════════════════════════════
logger.info("\nFINAL: Centering all images...")
for p in doc.paragraphs:
    if has_image(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)


# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
logger.info("\nSaving...")
doc.save(DOC_PATH)
logger.info(f"Saved to {DOC_PATH}")
logger.info("All fixes applied!")
