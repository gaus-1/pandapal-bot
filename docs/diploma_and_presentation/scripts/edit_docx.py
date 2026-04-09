import docx
import re
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
output_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"

doc = docx.Document(doc_path)

# 1. & 2. Headers and Titles
for i, p in enumerate(doc.paragraphs):
    if "АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ" in p.text and i < 5:
        # We replace the first 4 lines with the new title
        doc.paragraphs[i].text = "АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ ДОПОЛНИТЕЛЬНОГО ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ «АКАДЕМИЯ ТОП»"
        if i+1 < len(doc.paragraphs): doc.paragraphs[i+1].text = ""
        if i+2 < len(doc.paragraphs): doc.paragraphs[i+2].text = ""
        if i+3 < len(doc.paragraphs): doc.paragraphs[i+3].text = ""

    if "ИНДИВИДУАЛЬНЫЙ ПРОЕКТ" in p.text:
        p.text = p.text.replace("ИНДИВИДУАЛЬНЫЙ ПРОЕКТ", "ДИПЛОМНАЯ РАБОТА")

# 3. Add captions to tables if they don't have one
for i, table in enumerate(doc.tables):
    if i == 3:
        new_p = docx.oxml.OxmlElement('w:p')
        table._element.addprevious(new_p)
        p = docx.text.paragraph.Paragraph(new_p, doc._body)
        p.text = "Таблица 4 — Описание полей базы данных (часть 1)"
        p.style = doc.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif i == 4:
        new_p = docx.oxml.OxmlElement('w:p')
        table._element.addprevious(new_p)
        p = docx.text.paragraph.Paragraph(new_p, doc._body)
        p.text = "Таблица 5 — Описание полей базы данных (часть 2)"
        p.style = doc.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif i == 5:
        new_p = docx.oxml.OxmlElement('w:p')
        table._element.addprevious(new_p)
        p = docx.text.paragraph.Paragraph(new_p, doc._body)
        p.text = "Таблица 6 — Описание полей базы данных (часть 3)"
        p.style = doc.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 4. Replace dashes with bullets
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith(("–", "—", "-", "•")):
        clean_text = re.sub(r'^[\-\—\–\•]\s*', '', p.text.lstrip())
        p.text = "• " + clean_text
        try:
            p.style = 'List Paragraph'
        except:
            pass
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = 0

# 5. Красные строки (First line indents)
for p in doc.paragraphs:
    # Skip titles, centered text, empty lines, and lists
    if "List" not in p.style.name and p.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT) and len(p.text.strip()) > 3 and not p.text.startswith("• "):
        # Skip Ascii art
        if any(c in p.text for c in "┌│└├─"):
            continue
        # Check standard normal text
        if "Normal" in p.style.name:
            p.paragraph_format.first_line_indent = Cm(1.25)

# 6. Схемы подписать
# We know exact lines for diagrams approximately. Let's just look at previous paragraph to confirm we are leaving the diagram.
in_diagram = False
schema_counter = 1
titles = {
    1: "Схема бизнес-процессов платформы",
    2: "Архитектура и потоки данных",
    3: "ER-диаграмма базы данных",
    4: "Схема развёртывания платформы"
}

for i, p in enumerate(doc.paragraphs):
    is_ascii = any(c in p.text for c in "┌│└├─")
    if is_ascii:
        in_diagram = True
    elif in_diagram and not is_ascii and len(p.text.strip()) > 0:
        # We just exited a diagram!
        in_diagram = False

        # Check if we already have a caption "Рисунок N" in the next few paragraphs or current paragraph
        # Sometimes there's an empty paragraph between diagram and text
        already_caption = False
        if "Рисунок" in p.text:
            already_caption = True

        if not already_caption:
            # We insert a caption BEFORE the currently scanned non-diagram paragraph
            new_p = docx.oxml.OxmlElement('w:p')
            p._element.addprevious(new_p)
            caption_p = docx.text.paragraph.Paragraph(new_p, doc._body)

            title = titles.get(schema_counter, "Схема")
            caption_p.text = f"Рисунок {schema_counter} — {title}"
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.style = doc.styles['Normal']

            schema_counter += 1

doc.save(output_path)
print(f"Saved optimized document to {output_path}")
