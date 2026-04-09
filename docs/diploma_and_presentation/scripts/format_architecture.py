import logging
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)

DOC_PATH = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(DOC_PATH)

# Fix 1: Architecture text block
# Find paragraph with "Асинхронный микромонолит"
for i, p in enumerate(doc.paragraphs):
    if 'Асинхронный микромонолит' in p.text:
        text = p.text
        if '\n' in text:
            logger.info("Found broken Architecture list. Fixing.")
            lines = text.split('\n')
            p.text = lines[0].strip() # The intro text
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            # Make sure it's regular font
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

            # Insert correctly formatted list below it!
            anchor = p
            for line in lines[1:]:
                line = line.strip()
                if not line:
                    continue

                # Check if it's a top-level numbered item
                if line[0].isdigit() and line[1] == '.':
                    bullet_style = False
                    indent = Cm(1.25)
                    text_content = line[2:].strip()
                    prefix = line[:2] + " "
                elif line.startswith('-') or line.startswith('–') or line.startswith('—'):
                    bullet_style = True
                    indent = Cm(2.5)
                    text_content = line[1:].strip()
                    prefix = "– "
                else:
                    bullet_style = False
                    indent = Cm(1.25)
                    text_content = line
                    prefix = ""

                # Create paragraph
                new_p_elem = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
                # append after anchor
                anchor._element.addnext(new_p_elem)

                new_p = docx.text.paragraph.Paragraph(new_p_elem, doc._body)
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.left_indent = indent
                new_p.paragraph_format.first_line_indent = Cm(-1.25) if prefix else Cm(0)
                new_p.paragraph_format.space_after = Pt(6)

                run = new_p.add_run(prefix + text_content)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

                anchor = new_p
        break

# Fix 2: PAYMENTS Table formatting
# Find caption "Таблица А.3"
for i, p in enumerate(doc.paragraphs):
    if 'Таблица А.3' in p.text and 'PAYMENTS' in p.text:
        logger.info("Found PAYMENTS table caption.")
        # Ensure it's correctly formatted according to GOST
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.font.bold = False
            r.font.italic = False

        # The table should be the next element after the paragraph mathematically,
        # or we find the table whose previous paragraph is this one.
        # But doc.tables is an independent list.
        # We can find which table starts after this paragraph in XML.
        break

# Simpler way to fix ALL tables in doc:
logger.info("Fixing all document tables for beautiful symmetric proportions...")
for t in doc.tables:
    t.style = 'Table Grid'
    # Autofit contents
    t.autofit = True
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.first_line_indent = Pt(0)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

doc.save(DOC_PATH)
logger.info("Successfully fixed the architecture text list and table formatting!")
