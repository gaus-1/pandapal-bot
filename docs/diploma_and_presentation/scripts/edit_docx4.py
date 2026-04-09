import docx
import re
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

in_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
out_path = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_final.docx"
doc = docx.Document(in_path)

# --- 1. Headers & Title ---
doc.paragraphs[0].text = "АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ ДОПОЛНИТЕЛЬНОГО ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ «АКАДЕМИЯ ТОП»"
# Clear the next 3 lines
for i in range(1, 4):
    doc.paragraphs[i].text = ""

for p in doc.paragraphs:
    if "ИНДИВИДУАЛЬНЫЙ ПРОЕКТ" in p.text:
        p.text = p.text.replace("ИНДИВИДУАЛЬНЫЙ ПРОЕКТ", "ДИПЛОМНАЯ РАБОТА")
    if "10 ORM-моделей" in p.text:
        p.text = p.text.replace("10 ORM-моделей", "14 ORM-моделей")

# --- 2. Table Captions 4, 5, 6 ---
for i, table in enumerate(doc.tables):
    if i in [3, 4, 5]: # Tables 4, 5, 6
        new_p = docx.oxml.OxmlElement('w:p')
        table._element.addprevious(new_p)
        p = docx.text.paragraph.Paragraph(new_p, doc._body)
        p.text = f"Таблица {i+1} — Описание полей базы данных (часть {i-2})"
        p.style = doc.styles['Normal']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 3. TOC Formatting ---
toc_pattern = re.compile(r'^(.*?)([\.\…\s]+)(\d+)$')
for p in doc.paragraphs:
    text_val = p.text.strip()
    if text_val and "ОГЛАВЛЕНИЕ" not in text_val.upper() and not "ВВЕДЕНИЕ" == text_val.upper():
        # Match "Title .... 5"
        match = toc_pattern.match(text_val)
        if match and len(match.group(2)) > 1:
            title = match.group(1).rstrip('.\… ')
            page = match.group(3)
            p.text = f"{title}\t{page}"

            tab_stops = p.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            p.paragraph_format.first_line_indent = 0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 4. Dash to Bullet & Standard Indent ---
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith(("–", "—", "-")) and len(text) > 3:
        clean_text = re.sub(r'^[\-\—\–\•]\s*', '', p.text.lstrip())
        p.text = "• " + clean_text
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = 0
    elif "List" not in p.style.name and p.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
        if not any(c in text for c in "┌│└├─") and len(text) > 3 and not text.startswith("• "):
            if "Normal" in p.style.name or not p.style.name:
                p.paragraph_format.first_line_indent = Cm(1.25)

# --- 5. Inject 4 Captions at EXACT locations ---
# We will insert "Рисунок X" after specific trigger strings
diagram_triggers = [
    ("Общий поток данных можно описать следующей схемой:", "Схема бизнес-процессов платформы"),
    ("Общая архитектура системы может быть представлена в виде следующей диаграммы:", "Архитектура и потоки данных"),
    ("База данных PandaPal содержит 14 ORM-моделей", "ER-диаграмма базы данных"),
    ("Приложение Б. Схема развёртывания платформы", "Схема развёртывания платформы")
]

for trigger, caption_title in diagram_triggers:
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if trigger in p.text:
            start_idx = i
            break

    if start_idx != -1:
        # Find the end of the ascii diagram
        end_idx = start_idx
        in_ascii = False
        for j in range(start_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text.strip()
            is_ascii = any(c in text for c in "┌│└├─")
            if is_ascii:
                in_ascii = True
            elif in_ascii and not is_ascii and len(text) > 0:
                # the diagram ended
                end_idx = j - 1
                break

        if in_ascii:
            # We found the end of the diagram
            new_p = docx.oxml.OxmlElement('w:p')
            # Insert AFTER the end of the diagram
            doc.paragraphs[end_idx]._element.addnext(new_p)
            caption_p = docx.text.paragraph.Paragraph(new_p, doc._body)

            fig_num = diagram_triggers.index((trigger, caption_title)) + 1
            caption_p.text = f"Рисунок {fig_num} — {caption_title}"
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.first_line_indent = 0
            # Mark it so we don't double process
            caption_p._fig_caption = True

# --- 6. ER Diagram Rewrite ---
er_start = -1
er_end = -1
for i, p in enumerate(doc.paragraphs):
    if "База данных PandaPal содержит 14 ORM-моделей" in p.text:
        for j in range(i+1, len(doc.paragraphs)):
            if any(c in doc.paragraphs[j].text for c in "┌│└├─"):
                er_start = j
                break
        break

if er_start != -1:
    for j in range(er_start, len(doc.paragraphs)):
        if not any(c in doc.paragraphs[j].text for c in "┌│└├─") and len(doc.paragraphs[j].text.strip()) > 0:
            er_end = j - 1
            break

if er_start != -1 and er_end != -1:
    # Delete old ER diagram
    for j in range(er_start, er_end + 1):
        p = doc.paragraphs[er_start]
        p._element.getparent().remove(p._element)

    new_er_diagram = """
                              ┌──────────────────┐
                              │      users       │
                              │──────────────────│
                              │ telegram_id (PK) │
                              └────────┬─────────┘
      ┌─────────────────┬──────────────┼──────────────┬─────────────────┐
      │                 │              │              │                 │
      ▼                 ▼              ▼              ▼                 ▼
┌───────────┐     ┌───────────┐  ┌───────────┐  ┌───────────┐     ┌───────────┐
│panda_pets │     │ referrals │  │chat_histor│  │ learning  │     │   games   │
│───────────│     │───────────│  │───────────│  │───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│  │user_id(FK)│  │user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘  └───────────┘  └───────────┘     └───────────┘
      │                 │              │              │                 │
      ▼                 ▼              ▼              ▼                 ▼
┌───────────┐     ┌───────────┐  ┌───────────┐  ┌───────────┐     ┌───────────┐
│ payments  │     │ payouts   │  │   daily   │  │ homework  │     │game_stats │
│───────────│     │───────────│  │───────────│  │───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│  │user_id(FK)│  │user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘  └───────────┘  └───────────┘     └───────────┘
      │                 │
      ▼                 ▼
┌───────────┐     ┌───────────┐
│ subs      │     │ analytics │
│───────────│     │───────────│
│user_id(FK)│     │user_id(FK)│
└───────────┘     └───────────┘
"""
    anchor = doc.paragraphs[er_start]
    for line in new_er_diagram.strip('\n').split('\n'):
        new_p = docx.oxml.OxmlElement('w:p')
        anchor._element.addprevious(new_p)
        para = docx.text.paragraph.Paragraph(new_p, doc._body)
        para.text = line
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = 0
        for run in para.runs:
            run.font.name = 'Courier New'

# --- 7. Eliminate Multiple Empty Lines ---
empty_count = 0
to_remove = []
for p in doc.paragraphs:
    if not p.text.strip():
        empty_count += 1
        if empty_count >= 2:
            to_remove.append(p)
    else:
        empty_count = 0

for p in to_remove:
    p._element.getparent().remove(p._element)

# end trimming
while doc.paragraphs and not doc.paragraphs[-1].text.strip():
    p = doc.paragraphs[-1]
    p._element.getparent().remove(p._element)

doc.save(out_path)
