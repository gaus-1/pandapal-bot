"""
Конвертация diploma_project.md → DOCX.
ГОСТ Р 7.0.110-2025 / ГОСТ 7.32-2017:
  - Times New Roman 14pt, 1.5 интервал
  - Поля: лево 30мм, право 15мм, верх/низ 20мм
  - Выравнивание по ширине, красная строка 1.25 см
  - Нумерация страниц внизу по центру (с титульного, но без номера на нём)
  - Оглавление с правым табулятором и точечным заполнителем
  - Никаких Heading-стилей (всё inline XML)
"""

import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============================================================
# Low-level XML helpers
# ============================================================


def xml_spacing(p, line=360, before=0, after=0):
    """line: 240=single, 360=1.5, 480=double (in 1/240 of a line)."""
    pPr = p._element.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))


def xml_indent(p, first_line_cm=None, left_cm=None, hanging_cm=None):
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line_cm is not None:
        ind.set(qn("w:firstLine"), str(int(first_line_cm * 567)))
    if left_cm is not None:
        ind.set(qn("w:left"), str(int(left_cm * 567)))
    if hanging_cm is not None:
        ind.set(qn("w:hanging"), str(int(hanging_cm * 567)))


def xml_jc(p, val="both"):
    """val: both=justify, center, left, right."""
    pPr = p._element.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), val)


def make_run(p, text, font="Times New Roman", size=14, bold=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor(0, 0, 0)
    # XML-level font override for Google Docs compatibility
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    rf.set(qn("w:cs"), font)
    return r


def add_inline_code(p, text, size=14, bold=False):
    for part in re.split(r"(`[^`]+`)", text):
        if part.startswith("`") and part.endswith("`"):
            make_run(p, part[1:-1], "Courier New", size - 2, bold)
        else:
            make_run(p, part, "Times New Roman", size, bold)


# ============================================================
# Page numbering (bottom center, skip first page)
# ============================================================


def add_page_numbers(doc):
    """Add page numbers in the footer, bottom-center. Title page has no number."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # skip title page

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    xml_jc(fp, "center")

    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run1 = fp.add_run()
    run1._element.append(fldChar1)
    make_run(fp, "", size=12)  # placeholder

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2 = fp.add_run()
    run2._element.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3 = fp.add_run()
    run3._element.append(fldChar2)

    # Style the page number runs
    for r in fp.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)

    # First page footer — empty (no number on title page)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].text = ""


# ============================================================
# TOC with proper right-aligned tab stop + dot leader
# ============================================================


def toc_line(doc, section_name, page_num, bold=False, indent=False):
    """TOC: visible dots + right tab stop for page number alignment (ГОСТ 7.32-2017)."""
    p = doc.add_paragraph()

    # 1) Right-aligned tab stop at right margin
    pPr = p._element.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    tab_el = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:leader"), "dot")
    tab_el.set(qn("w:pos"), str(int(16.5 * 567)))
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    # 2) Section name
    make_run(p, section_name, bold=bold)

    # 3) Visible dots (conservative count to not overflow past tab stop)
    # TNR 14pt avg char width ~2.4mm; available width 165mm
    avail_mm = 165 - (7.5 if indent else 0)
    used_mm = len(section_name) * 2.4
    dot_width_mm = 1.8  # dot character is narrower than letters
    ndots = max(3, int((avail_mm - used_mm - 10) / dot_width_mm))
    make_run(p, " " + "." * ndots)

    # 4) Tab to right margin + page number (tab aligns at exact position)
    r = p.add_run("\t" + page_num)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)
    rPr = r._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:cs"), "Times New Roman")
    rPr.insert(0, rf)

    xml_spacing(p, 360)
    if indent:
        xml_indent(p, left_cm=0.75)
    return p


# ============================================================
# Paragraph constructors
# ============================================================


def body(doc, text, bold=False, indent=True, center=False, size=14):
    p = doc.add_paragraph()
    add_inline_code(p, text, size, bold)
    xml_jc(p, "center" if center else "both")
    xml_spacing(p, 360)
    if indent and not center:
        xml_indent(p, first_line_cm=1.25)
    return p


def heading(doc, text, before=240, after=120):
    p = doc.add_paragraph()
    make_run(p, text, bold=True)
    xml_jc(p, "center")
    xml_spacing(p, 360, before, after)
    return p


def section_h(doc, text):
    p = doc.add_paragraph()
    make_run(p, text, bold=True)
    xml_jc(p, "both")
    xml_spacing(p, 360, 240, 120)
    xml_indent(p, first_line_cm=1.25)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    prefix = "–  " if level == 0 else "•  "
    add_inline_code(p, prefix + text)
    xml_jc(p, "both")
    xml_spacing(p, 360)
    xml_indent(p, left_cm=1.25 + level * 0.75, hanging_cm=0.5)
    return p


def pgbrk(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)
    # Zero spacing so it doesn't add blank space
    xml_spacing(p, 0, 0, 0)


def add_table(doc, rows):
    if not rows:
        return
    nc = max(len(r.split("\t")) for r in rows)
    t = doc.add_table(rows=len(rows), cols=nc)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = row.split("\t")
        for j, ct in enumerate(cells):
            if j < nc:
                c = t.rows[i].cells[j]
                c.text = ""
                cp = c.paragraphs[0]
                make_run(cp, ct.strip(), size=12, bold=(i == 0))
                xml_spacing(cp, 240, 0, 0)


def add_code(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        make_run(p, ln, "Courier New", 9)
        xml_spacing(p, 240, 0, 0)
        xml_indent(p, left_cm=1.0)


# ============================================================
# Classification
# ============================================================

H1 = {
    "СОДЕРЖАНИЕ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ",
    "ПАСПОРТ ПРОЕКТА",
    "ПРИЛОЖЕНИЯ",
}
TITLE = {
    "АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ",
    "ПРОФЕССИОНАЛЬНАЯ ОБРАЗОВАТЕЛЬНАЯ ОРГАНИЗАЦИЯ",
    "МОСКОВСКИЙ МЕЖДУНАРОДНЫЙ КОЛЛЕДЖ ЦИФРОВЫХ",
    "ТЕХНОЛОГИЙ «АКАДЕМИЯ ТОП»",
    "ИНДИВИДУАЛЬНЫЙ ПРОЕКТ",
}
SUBS = {
    "Актуальность темы",
    "Проблема проекта",
    "Объект и предмет исследования",
    "Цель проекта",
    "Задачи проекта",
    "Планируемый проектный продукт",
    "Целевая аудитория проекта",
    "Ресурсное обеспечение проекта",
    "Ожидаемые результаты",
    "Возможные риски и пути их преодоления",
    "Методы исследования",
    "Научная новизна",
    "Практическая значимость",
    "Нормативные правовые акты",
    "Книги",
    "Статьи из периодических изданий и сборников",
    "Иностранные источники",
    "Электронные ресурсы",
    "Серверная часть (Backend):",
    "Клиентская часть (Frontend):",
    "Внешние интеграции:",
    "Инфраструктура:",
}
SUBS_STRIPPED = {k.rstrip(":") for k in SUBS}


# ============================================================
# Main processor
# ============================================================


def process(doc, path):
    with open(path, encoding="utf-8") as f:
        lines = [line_str.rstrip("\r\n") for line_str in f]

    # --- Pre-scan: ПАСПОРТ ---
    ps, pe = None, None
    for idx, ln in enumerate(lines):
        if ln.strip() == "ПАСПОРТ ПРОЕКТА":
            ps = idx + 1
        if ps and not pe and ln.strip() == "________________" and idx > ps + 3:
            pe = idx
            break
    passport_rows, passport_set = [], set()
    if ps and pe:
        for idx in range(ps, pe):
            if "\t" in lines[idx]:
                passport_rows.append(lines[idx])
            passport_set.add(idx)

    # --- Pre-scan: TOC ---
    ts, te = None, None
    for idx, ln in enumerate(lines):
        if ln.strip() == "СОДЕРЖАНИЕ":
            ts = idx + 1
        if ts and not te and ln.strip() == "________________" and idx > ts + 3:
            te = idx
            break
    toc_set = set()
    if ts and te:
        for idx in range(ts, te):
            toc_set.add(idx)

    i = 0
    in_code = False
    code_buf = []

    # Find first page break line for title page detection
    first_pb = None
    for idx, ln in enumerate(lines):
        if ln.strip() == "________________":
            first_pb = idx
            break

    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # --- Code block ---
        if s.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code(doc, code_buf)
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # --- Page break (no extra empty paragraph) ---
        if s == "________________":
            pgbrk(doc)
            i += 1
            continue

        # --- ПАСПОРТ ПРОЕКТА ---
        if s == "ПАСПОРТ ПРОЕКТА":
            heading(doc, "ПАСПОРТ ПРОЕКТА")
            for row in passport_rows:
                parts = row.split("\t")
                label = parts[0].strip() if parts else ""
                value = parts[1].strip() if len(parts) > 1 else ""
                if label:
                    p = doc.add_paragraph()
                    make_run(p, label, bold=True)
                    xml_spacing(p, 360)
                    xml_indent(p, first_line_cm=1.25)
                if value:
                    p = doc.add_paragraph()
                    make_run(p, value)
                    xml_jc(p, "both")
                    xml_spacing(p, 360)
                    xml_indent(p, first_line_cm=2.0)
            i = pe if pe else i + 1
            continue

        if i in passport_set:
            i += 1
            continue

        # --- TOC entries ---
        if i in toc_set:
            if "\t" in line:
                parts = line.split("\t")
                sec = parts[0].strip()
                pg = parts[1].strip() if len(parts) > 1 else ""
                is_bold = sec.startswith("ГЛАВА") or sec in H1
                is_subsection = re.match(r"^\d+\.\d+\.", sec)
                toc_line(doc, sec, pg, bold=is_bold, indent=bool(is_subsection))
            i += 1
            continue

        # --- Empty line ---
        if s == "":
            # Preserve empty lines on title page for vertical spacing
            if first_pb and i < first_pb:
                p = doc.add_paragraph()
                make_run(p, "")
                xml_spacing(p, 360)
            i += 1
            continue

        # --- Tab data (real tables) ---
        if "\t" in line and s not in TITLE:
            rows = []
            while i < len(lines):
                cur = lines[i].strip()
                if "\t" in lines[i] and cur and cur != "________________":
                    rows.append(lines[i])
                    i += 1
                elif cur == "":
                    pk = i + 1
                    while pk < len(lines) and lines[pk].strip() == "":
                        pk += 1
                    if pk < len(lines) and "\t" in lines[pk]:
                        i += 1
                        continue
                    else:
                        i += 1
                        break
                else:
                    break
            add_table(doc, rows)
            continue

        # --- ГЛАВА heading ---
        if re.match(r"^ГЛАВА \d+\.", s):
            heading(doc, s)
            i += 1
            continue

        # --- H1 keywords ---
        if s in H1:
            heading(doc, s)
            i += 1
            continue

        # --- Section heading (1.1., 2.3.) ---
        if re.match(r"^\d+\.\d+\.\s+", s):
            section_h(doc, s)
            i += 1
            continue

        # --- Appendix heading ---
        if re.match(r"^Приложение [А-Я]\.\s+", s):
            section_h(doc, s)
            i += 1
            continue

        # --- Title page centered ---
        if s in TITLE:
            body(doc, s, bold=True, center=True, indent=False)
            i += 1
            continue

        # --- Title fields ---
        if (
            s.startswith("Тема:")
            or s.startswith("Тип проекта:")
            or s.startswith("Предметная область:")
        ):
            p = doc.add_paragraph()
            c = s.index(":")
            make_run(p, s[: c + 1], bold=True)
            make_run(p, s[c + 1 :])
            xml_jc(p, "center")
            xml_spacing(p, 360)
            i += 1
            continue

        if s.startswith("Автор:") or s.startswith("Руководитель:"):
            p = doc.add_paragraph()
            c = s.index(":")
            make_run(p, s[: c + 1], bold=True)
            make_run(p, s[c + 1 :])
            xml_spacing(p, 360)
            i += 1
            continue

        if re.match(r"^Санкт-Петербург", s):
            body(doc, s, center=True, indent=False)
            i += 1
            continue

        # --- Subheadings (bold body paragraph) ---
        if s in SUBS or s.rstrip(":") in SUBS_STRIPPED:
            body(doc, s, bold=True, indent=True)
            i += 1
            continue

        # --- «Сущее» / «Должное» ---
        if s.startswith("«Сущее»") or s.startswith("«Должное»"):
            body(doc, s, bold=True, indent=True)
            i += 1
            continue

        # --- Table caption ---
        if s.startswith("Таблица ") and "—" in s:
            body(doc, s, bold=True, indent=False)
            i += 1
            continue

        # --- Bullets ---
        if s.startswith("* "):
            bullet(doc, s[2:], level=0)
            i += 1
            continue

        if line.startswith("   * ") or line.startswith("   * "):
            text = s.lstrip("*").lstrip("*").strip()
            bullet(doc, text, level=1)
            i += 1
            continue

        # --- Normal paragraph ---
        body(doc, s, indent=True)
        i += 1

    return doc


def create_doc():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)

    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    st.font.color.rgb = RGBColor(0, 0, 0)
    pf = st.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    return doc


if __name__ == "__main__":
    doc = create_doc()
    process(doc, "diploma_project.md")
    add_page_numbers(doc)
    out = "diplom_project(PandaPal).docx"
    doc.save(out)
    print(f"OK: {out}")
