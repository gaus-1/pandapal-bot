from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading_center(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table_erd(doc, title, rows):
    doc.add_paragraph(title).runs[0].font.bold = True
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Поле"
    hdr_cells[1].text = "Тип данных"
    hdr_cells[2].text = "Описание"

    for row in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]

    doc.add_paragraph("")  # spacing


def main():
    target_docx = (
        r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal).docx"
    )
    doc = Document(target_docx)

    # We will find the paragraph saying "ПРИЛОЖЕНИЕ А" and delete everything from there to the end.
    app_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "ПРИЛОЖЕНИЕ А" in p.text:
            app_index = i
            break

    if app_index != -1:
        # Delete paragraphs from app_index to end
        for i in range(len(doc.paragraphs) - 1, app_index - 1, -1):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)

    # Now append clean native text/tables!
    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ А")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Схема базы данных (ERD)")

    # Table 1
    add_table_erd(
        doc,
        "Таблица А.1 — Структура сущности USERS (Пользователи)",
        [
            ("id (PK)", "Integer", "Уникальный идентификатор"),
            ("telegram_id", "BigInt", "ID пользователя в Telegram"),
            ("current_state", "String", "Текущее состояние в FSM"),
            ("is_premium", "Boolean", "Флаг платной подписки"),
            ("total_coins", "Integer", "Баланс виртуальной валюты"),
        ],
    )

    # Table 2
    add_table_erd(
        doc,
        "Таблица А.2 — Структура сущности PANDA_PET (Тамагочи)",
        [
            ("id (PK)", "Integer", "Уникальный идентификатор"),
            ("user_id (FK)", "Integer", "Связь с таблицей USERS"),
            ("health", "Integer", "Уровень здоровья (0-100)"),
            ("energy", "Integer", "Уровень энергии (0-100)"),
            ("satiety", "Integer", "Уровень сытости (0-100)"),
            ("mood", "Integer", "Уровень настроения (0-100)"),
        ],
    )

    # Table 3
    add_table_erd(
        doc,
        "Таблица А.3 — Структура сущности PAYMENTS (Платежи)",
        [
            ("id (PK)", "Integer", "Уникальный идентификатор"),
            ("user_id (FK)", "Integer", "Связь с таблицей USERS"),
            ("yookassa_id", "String", "Внешний ID транзакции ЮKassa"),
            ("amount", "Decimal", "Сумма платежа"),
            ("status", "String", "Статус (pending, succeeded, canceled)"),
        ],
    )

    doc.add_page_break()
    add_heading_center(doc, "ПРИЛОЖЕНИЕ Б")
    add_heading_center(doc, "(обязательное)")
    add_heading_center(doc, "Архитектура программного комплекса")

    p = doc.add_paragraph()
    r = p.add_run(
        "Архитектура системы построена по паттерну «Асинхронный микромонолит» и состоит из следующих слоев:\n\n"
        "1. КЛИЕНТСКИЙ СЛОЙ (Frontend)\n"
        "   — Telegram Bot (aiogram 3): мессенджер-интерфейс\n"
        "   — Telegram Mini App (React 19, Zustand): встроенное веб-приложение\n"
        "   — Landing Page (React, Tailwind CSS): промо-сайт\n\n"
        "2. СЕРВЕРНЫЙ СЛОЙ (Backend)\n"
        "   — Entrypoint: aiohttp web-сервер (обработка Webhook)\n"
        "   — Security / Middleware: 4-уровневый AI модератор (RegExp, Levenshtein, ML-Classifier)\n"
        "   — Business Logic Services: Game Engine, RAG Engine, Payments API\n\n"
        "3. СЛОЙ ДАННЫХ (Databases)\n"
        "   — Основная БД: PostgreSQL 17 + расширение pgvector (индексы HNSW)\n"
        "   — Быстрый кэш и стейты: Redis 7.1 (In-Memory)\n\n"
        "4. ВНЕШНИЕ ИНТЕГРАЦИИ (External APIs)\n"
        "   — Yandex Cloud (SpeechKit STT, Vision OCR, YandexGPT Pro)\n"
        "   — YooKassa API (вебхуки для обработки платежей)"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    fixed_docx = r"c:\Users\Vyacheslav\PandaPal\docs\diploma_and_presentation\diplom_project(PandaPal)_fixed.docx"
    doc.save(fixed_docx)
    print("Cleaned and injected native tables into _fixed.docx!")


if __name__ == "__main__":
    main()
